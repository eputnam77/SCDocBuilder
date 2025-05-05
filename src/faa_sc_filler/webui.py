import logging
import gradio as gr
from typing import Tuple, Optional, Dict
import re
from pathlib import Path
import tempfile
import shutil
from docx import Document
from .processor import DocumentProcessor

# Configure logging to show debug messages
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def create_temp_dir():
    """Create temporary directory for file uploads"""
    temp_dir = Path("temp_uploads")
    temp_dir.mkdir(exist_ok=True)
    return temp_dir

def save_uploaded_file(file_data, filename: str) -> Path:
    """Save uploaded file data to temporary file"""
    temp_dir = create_temp_dir()
    temp_file = temp_dir / filename
    
    if isinstance(file_data, (str, Path)):
        # File is already a path
        return Path(file_data)
        
    if hasattr(file_data, 'name'):
        # File is a file-like object
        return Path(file_data.name)
        
    # File is binary data
    with open(temp_file, 'wb') as f:
        if isinstance(file_data, bytes):
            f.write(file_data)
        else:
            shutil.copyfileobj(file_data, f)
    return temp_file

def validate_file(file: str) -> Tuple[bool, str]:
    """Validate uploaded file path."""
    if not file:
        return False, "No file uploaded"
    path = Path(file)
    if not path.exists():
        return False, f"File not found: {file}"
    if not path.suffix.lower() == '.docx':
        return False, f"Invalid file type: {path.suffix}"
    return True, ""

def generate(template: str, worksheet: str, dry_run: bool = False) -> Tuple[Optional[str], Optional[dict]]:
    """Process document through web interface."""
    try:
        logger.info("Starting document generation")
        logger.debug(f"Template file: {template}")
        logger.debug(f"Worksheet file: {worksheet}")

        # Validate files
        template_valid, template_error = validate_file(template)
        worksheet_valid, worksheet_error = validate_file(worksheet)

        if not template_valid:
            return None, {"error": f"Template file error: {template_error}"}
        if not worksheet_valid:
            return None, {"error": f"Worksheet file error: {worksheet_error}"}
        
        processor = DocumentProcessor()
        
        # Load documents from file paths
        template_doc = Document(template)
        worksheet_doc = Document(worksheet)
        
        # Extract data first
        worksheet_data = processor.extractor.extract_data(worksheet_doc)
        logger.debug(f"Extracted data: {worksheet_data}")
        
        # Process document with temporary output path
        output_path = "output/processed_SC.docx" if not dry_run else None
        path, diff = processor.process_document(
            template=template_doc,
            replacements=worksheet_data,
            output_path=output_path,
            dry_run=dry_run
        )
        
        logger.info("Document processing complete")
        return path, diff if dry_run else None
        
    except Exception as e:
        logger.exception("Error in document generation")
        return None, {"error": str(e)}

def validate_inputs(cfr_part: str, docket_no: str, notice_no: str) -> Dict[str, str]:
    """Validate user inputs and return error messages."""
    errors = {}
    valid_parts = {"23", "25", "27", "29", "31", "33", "35"}
    
    if cfr_part:
        parts = {p.strip() for p in cfr_part.split(",")}
        if not all(p in valid_parts for p in parts):
            errors["cfr_part"] = "Invalid CFR Part(s). Use comma-separated values from: 23,25,27,29,31,33,35"
            
    if docket_no and not re.match(r'^FAA-\d{4}-\d{4}$', docket_no):
        errors["docket_no"] = "Invalid format. Use FAA-YYYY-XXXX"
        
    if notice_no and not re.match(r'^\d{2}-\d{2}-\d{2}-SC$', notice_no):
        errors["notice_no"] = "Invalid format. Use XX-XX-XX-SC"
        
    return errors

def enhanced_generate(template_file, worksheet_file, cfr_part=None, docket_no=None, notice_no=None, dry_run=False):
    """Enhanced document generation with optional fields."""
    try:
        # Load documents first
        template_doc = Document(template_file)
        worksheet_doc = Document(worksheet_file)
        
        # Process document with optional fields
        processor = DocumentProcessor()
        worksheet_data = processor.extractor.extract_data(worksheet_doc)
        
        # Add optional fields if provided
        if cfr_part: worksheet_data["{CFRPart}"] = cfr_part
        if docket_no: worksheet_data["{DocketNo}"] = docket_no
        if notice_no: worksheet_data["{NoticeNo}"] = notice_no
        
        return processor.process_document(template_doc, worksheet_data, dry_run=dry_run)
        
    except Exception as e:
        logger.exception("Error in enhanced document generation")
        return None, {"error": str(e)}

def handle_generation(*args):
    template, worksheet, cfr, docket, notice, dry_run = args
    try:
        if not template or not worksheet:
            return None, {"error": "Both template and worksheet files are required"}, "Both files are required"

        # Save and process files
        temp_dir = Path("temp_uploads")
        temp_dir.mkdir(exist_ok=True)
        temp_template = temp_dir / "temp_template.docx"
        temp_worksheet = temp_dir / "temp_worksheet.docx"
        
        with open(temp_template, "wb") as f:
            f.write(template)
        with open(temp_worksheet, "wb") as f:
            f.write(worksheet)

        try:
            template_doc = Document(temp_template)
            worksheet_doc = Document(temp_worksheet)
            
            processor = DocumentProcessor()
            worksheet_data = processor.extractor.extract_data(worksheet_doc)
            
            if cfr: worksheet_data["{CFRPart}"] = cfr
            if docket: worksheet_data["{DocketNo}"] = docket
            if notice: worksheet_data["{NoticeNo}"] = notice
            
            # Process document with absolute path
            output_dir = Path.cwd() / "output"
            output_dir.mkdir(exist_ok=True)
            output_path = output_dir / "processed_SC.docx"
            
            path, diff = processor.process_document(
                template=template_doc,
                replacements=worksheet_data,
                output_path=str(output_path) if not dry_run else None,
                dry_run=True  # Always generate diff for preview
            )
            
            logger.info(f"Generated file path: {path}")
            logger.info(f"Generated diff: {diff}")
            
            if dry_run:
                return None, diff, "Showing preview (dry-run)"
            else:
                # Process again without dry_run to generate file
                path, _ = processor.process_document(
                    template=template_doc,
                    replacements=worksheet_data,
                    output_path=str(output_path),
                    dry_run=False
                )
                return str(output_path), diff, "Document generated successfully!"

        finally:
            temp_template.unlink(missing_ok=True)
            temp_worksheet.unlink(missing_ok=True)

    except Exception as e:
        logger.exception("Error in document generation")
        return None, {"error": str(e)}, f"Error: {str(e)}"

def create_ui() -> gr.Blocks:
    """Create the Gradio web interface."""
    output_dir = Path.cwd() / "output"
    output_dir.mkdir(exist_ok=True)
    
    with gr.Blocks(title="FAA Special Conditions Template Filler") as demo:
        gr.Markdown("### File Upload Debug Info")
        debug_output = gr.Markdown()
        
        template_file = gr.File(
            label="SC Template (.docx)",
            file_types=['.docx'],
            type="binary"
        )
        worksheet_file = gr.File(
            label="Worksheet (.docx)",
            file_types=['.docx'],
            type="binary"
        )

        def update_debug_info(template, worksheet):
            """Show readable debug info for uploaded files"""
            debug_info = []
            if template:
                debug_info.append(f"Template: {type(template).__name__}")
                if hasattr(template, 'name'):
                    debug_info.append(f"Template path: {template.name}")
            if worksheet:
                debug_info.append(f"Worksheet: {type(worksheet).__name__}")
                if hasattr(worksheet, 'name'):
                    debug_info.append(f"Worksheet path: {worksheet.name}")
            return "\n".join(debug_info) if debug_info else "No files uploaded yet"

        # Add file change handlers
        template_file.change(
            fn=update_debug_info,
            inputs=[template_file, worksheet_file],
            outputs=[debug_output]
        )
        worksheet_file.change(
            fn=update_debug_info,
            inputs=[template_file, worksheet_file],
            outputs=[debug_output]
        )
        
        gr.Markdown("### Optional Fields")
        with gr.Row():
            cfr_part = gr.Textbox(label="CFR Part(s) (comma-separated)", placeholder="25,27")
            docket_no = gr.Textbox(label="Docket No.", placeholder="FAA-2024-0001")
            notice_no = gr.Textbox(label="Notice No.", placeholder="24-01-01-SC")
            
        dry_run = gr.Checkbox(label="Dry-run (preview JSON diff)")
        
        with gr.Row():
            with gr.Column(scale=1):
                output_status = gr.Markdown("No document generated yet")
                output_view = gr.File(
                    label="Generated Document",
                    type="file",
                    file_count="single",
                    interactive=False,
                    visible=True
                )
            with gr.Column(scale=1):
                diff_view = gr.JSON(
                    label="Changes Preview",
                    container=True,
                    show_label=True
                )
        
        error_view = gr.Markdown()

        generate_btn = gr.Button("Generate", variant="primary")
        generate_btn.click(
            fn=handle_generation,
            inputs=[template_file, worksheet_file, cfr_part, docket_no, notice_no, dry_run],
            outputs=[output_view, diff_view, output_status]
        )

        gr.Markdown("""
        ## Instructions
        1. Upload template DOCX file
        2. Upload worksheet DOCX file
        3. Optionally enter CFR Part, Docket No., and Notice No.
        4. Click Generate
        """)
    return demo

def main():
    demo = create_ui()
    demo.queue()  # Enable queueing
    demo.launch(share=True)

if __name__ == "__main__":
    main()
