import logging
from typing import Dict, Optional, Tuple, Union

from docx.document import Document as DocumentType
from .config import DEFAULT_CONFIG, EXIT_SUCCESS, EXIT_REPLACEMENT_ERROR
from .replacer import PlaceholderReplacer
from .extractor import WorksheetExtractor
from .validator import DocumentValidator

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Orchestrates document processing operations."""

    def __init__(self, need_token: str = DEFAULT_CONFIG["need_token"]):
        logger.debug("Initializing DocumentProcessor")
        self.extractor = WorksheetExtractor()
        self.replacer = PlaceholderReplacer(need_token=need_token)
        self.validator = DocumentValidator()
        self.replacements = {}

    def process_sections(self, doc: DocumentType) -> None:
        """Process header and footer placeholders."""
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                self.replacer.process_paragraph(paragraph, self.replacements)
            for paragraph in section.footer.paragraphs:
                self.replacer.process_paragraph(paragraph, self.replacements)

    def process_conditional_blocks(
        self, doc: DocumentType, selected_option: int
    ) -> None:
        """Keep only the selected conditional block."""
        option_tag = f"OPTION_{selected_option}"
        for paragraph in list(doc.paragraphs):
            import re

            match = re.match(r"\[\[(OPTION_\d+)\]\](.*)\[\[/\1\]\]", paragraph.text)
            if not match:
                continue
            if match.group(1) == option_tag:
                paragraph.text = match.group(2)
            else:
                paragraph._element.getparent().remove(paragraph._element)

    def process_document(
        self,
        template: Union[str, DocumentType],
        replacements: Dict[str, str],
        output_path: Optional[str] = None,
        dry_run: bool = False,
    ) -> Tuple[Optional[str], Optional[Dict]]:
        """Process document with replacements."""
        logger.debug(f"Starting document processing with dry_run={dry_run}")
        logger.debug(f"Replacements: {replacements}")

        if dry_run:
            # Build diff dictionary for dry run
            diff = {}
            for key, value in replacements.items():
                logger.debug(f"Adding diff entry: {key} -> {value}")
                diff[key] = {"old": key, "new": value}
            return None, diff

        try:
            if isinstance(template, str):
                self.validator.validate_docx(template)
            # Use PlaceholderReplacer's process_document directly
            self.replacer.process_document(
                template,
                replacements,
                output_path=output_path,
            )

            return output_path, None

        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            raise


# Keep the standalone process_document function for backwards compatibility
def process_document(
    template_path: str, replacements: Dict[str, str], output_path: str
) -> int:
    """Process document replacing placeholders with values."""
    logger.debug("Starting document processing...")
    try:
        replacer = PlaceholderReplacer()
        replacer.process_document(template_path, replacements, output_path)
        return EXIT_SUCCESS
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}")
        return EXIT_REPLACEMENT_ERROR
