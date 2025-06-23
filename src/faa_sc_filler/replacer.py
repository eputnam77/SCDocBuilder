import logging
from typing import Dict, List, Optional, Union
from docx.text.paragraph import Paragraph
from docx import Document

logger = logging.getLogger(__name__)


class PlaceholderReplacer:
    """Handles replacement of placeholders in Word documents."""

    def __init__(self):
        logger.debug("Initializing PlaceholderReplacer")

    def find_placeholders_in_paragraph(self, paragraph: Paragraph) -> List[str]:
        """Find all placeholders in a paragraph."""
        logger.debug(f"Finding placeholders in paragraph: '{paragraph.text[:50]}...'")

        text = paragraph.text
        # Look for text between curly braces
        import re

        placeholders = re.findall(r"\{[^}]+\}", text)

        logger.debug(f"Found placeholders: {placeholders}")
        return placeholders

    def join_split_runs(self, paragraph: Paragraph) -> None:
        """Join runs that might contain split placeholders."""
        logger.debug(f"Joining split runs in paragraph: '{paragraph.text[:50]}...'")
        logger.debug(f"Number of runs before joining: {len(paragraph.runs)}")

        if len(paragraph.runs) <= 1:
            logger.debug("No runs to join")
            return

        # Collect all text from runs
        full_text = "".join(run.text for run in paragraph.runs)

        # Clear all runs except the first one
        first_run = paragraph.runs[0]
        first_run.text = full_text

        for run in paragraph.runs[1:]:
            run.text = ""

        logger.debug(f"Joined text: '{full_text[:50]}...'")
        logger.debug(f"Number of runs after joining: {len(paragraph.runs)}")

    def process_paragraph(
        self, paragraph: Paragraph, replacements: Dict[str, str]
    ) -> None:
        """Process a paragraph for placeholder replacements."""
        logger.debug(f"Processing paragraph: '{paragraph.text[:50]}...'")

        # 0. Skip empty paragraphs or those without placeholder markers
        if not paragraph.text or "{" not in paragraph.text or "}" not in paragraph.text:
            logger.debug("No placeholder markers found, skipping")
            return

        # 1. Find actual placeholders
        placeholders = self.find_placeholders_in_paragraph(paragraph)
        if not placeholders:
            logger.debug("No valid placeholders found, skipping")
            return

        logger.debug(f"Found placeholders to process: {placeholders}")

        # 2. Join runs and replace
        full_text = "".join(run.text for run in paragraph.runs)
        modified = False

        # Process each placeholder found in the text
        for placeholder in placeholders:
            placeholder_clean = placeholder.strip("{}")
            if (
                placeholder_clean not in replacements
                or not replacements[placeholder_clean].strip()
            ):
                logger.debug(
                    f"Missing replacement for '{placeholder_clean}', using NEED"
                )
                full_text = full_text.replace(placeholder, f"NEED: {placeholder_clean}")
            else:
                logger.debug(
                    f"Replacing '{placeholder}' with '{replacements[placeholder_clean]}'"
                )
                full_text = full_text.replace(
                    placeholder, replacements[placeholder_clean]
                )
            modified = True

        if modified:
            logger.debug(f"Writing modified text: '{full_text[:50]}...'")
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ""

    def process_document(
        self,
        template: Union[str, Document],
        replacements: Dict[str, str],
        output_path: Optional[str] = None,
    ) -> Document:
        """Process the entire document for replacements."""
        logger.debug(f"Processing document: {template}")
        logger.debug(f"Output path: {output_path}")
        logger.debug(f"Replacements to apply: {replacements}")
        logger.debug("Final replacements dict: %r", replacements)

        from pathlib import Path

        try:
            if isinstance(template, str):
                template_file = Path(template)
                if not template_file.exists():
                    raise FileNotFoundError(f"Template file not found: {template}")
                logger.debug("Template file exists: %s", template_file.absolute())
                logger.debug(
                    "Template file size: %d bytes",
                    template_file.stat().st_size,
                )
                doc = Document(template)
            else:
                doc = template

            # Dump first few paragraphs to verify content
            for i, p in enumerate(doc.paragraphs[:5]):
                logger.debug("Paragraph %d raw text: %r", i, p.text)
                # Also dump runs to see if placeholders are split
                if len(p.runs) > 1:
                    for j, run in enumerate(p.runs):
                        logger.debug("  Run %d.%d text: %r", i, j, run.text)

            # Quick check of content
            sample_text = doc.paragraphs[0].text if doc.paragraphs else "NO TEXT"
            logger.debug("First paragraph preview: '%s...'", sample_text[:100])
            logger.debug("Document loaded successfully")

            # Process all paragraphs
            logger.debug("Processing %d paragraphs", len(doc.paragraphs))
            for paragraph in doc.paragraphs:
                self.process_paragraph(paragraph, replacements)

            # Process all tables
            logger.debug("Processing %d tables", len(doc.tables))
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self.process_paragraph(paragraph, replacements)

            # Process headers and footers in all sections
            logger.debug("Processing headers and footers")
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        self.process_paragraph(paragraph, replacements)
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        self.process_paragraph(paragraph, replacements)

            if output_path:
                logger.debug("Saving document to: %s", output_path)
                doc.save(output_path)
            logger.info("Document processing completed successfully")

            return doc

        except Exception as e:
            logger.error("Error processing document: %s", str(e))
            raise
