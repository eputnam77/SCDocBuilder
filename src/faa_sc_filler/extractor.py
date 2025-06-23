import re
import logging
from typing import Dict, Optional, List, Any
from docx.table import Table
from docx import Document
from .config import FIELD_MAPPINGS, MULTILINE_FIELDS, CHECKBOX_MAPPINGS, _normalize

logger = logging.getLogger(__name__)


class WorksheetExtractor:
    def __init__(self):
        logger.debug("Initializing WorksheetExtractor")
        self.field_mappings = FIELD_MAPPINGS
        self.multiline_fields = MULTILINE_FIELDS
        self.checkbox_mappings = CHECKBOX_MAPPINGS
        self.cfr_pattern = re.compile(r"(?:14\s*CFR\s*)?[Pp]art\s*(\d+)", re.IGNORECASE)
        logger.debug("Initialized CFR pattern for extraction")

    def clean_text(self, text: str) -> str:
        """Remove extra whitespace, field labels, and normalize format."""
        logger.debug(f"Cleaning text: '{text}'")
        # Remove leading/trailing whitespace and normalize internal spaces
        text = " ".join(text.strip().split())
        # Remove enumeration (a., b., 1., etc)
        text = re.sub(r"^[a-z0-9]\.\s*", "", text, flags=re.IGNORECASE)
        # Normalize multiple colons
        text = re.sub(r":\s*:", ":", text)
        # Normalize spaces around colons
        text = re.sub(r"\s*:\s*", ":", text)
        logger.debug(f"Cleaned text result: '{text}'")
        return text

    def extract_field_value(self, text: str, field: str) -> Optional[str]:
        """Extract field value handling various formats."""
        logger.debug(f"Extracting field value for field '{field}' from text: '{text}'")
        if not text.strip():
            logger.debug("Empty text provided, returning None")
            return None

        # Strip all spaces around colons in both text and field
        input_text = re.sub(r"\s*:\s*", ":", text.strip())
        logger.debug(f"Normalized input text: '{input_text}'")
        base_field = re.sub(r"\s*:\s*", ":", field.rstrip(":").strip())

        # Remove any enumeration prefixes for comparison
        cleaned_text = re.sub(r"^[a-z0-9]\.\s*", "", input_text, flags=re.IGNORECASE)
        cleaned_field = re.sub(r"^[a-z0-9]\.\s*", "", base_field, flags=re.IGNORECASE)

        # Get field name without the trailing colon
        field_name = cleaned_field.rstrip(":")

        # Try to find the field and extract value
        if f"{field_name}:" in cleaned_text:
            # Split on the field name with colon
            _, value = cleaned_text.split(f"{field_name}:", 1)
            logger.debug(f"Found value with colon: '{value.strip()}'")
            return value.strip()

        # Try without colon
        if field_name in cleaned_text:
            _, value = cleaned_text.split(field_name, 1)
            # Clean up any remaining colons
            return value.lstrip(":").strip()

        # Special handling for dates in text
        if "scheduled for" in cleaned_text and any(
            d in field_name.lower() for d in ["date", "certification"]
        ):
            match = re.search(r"scheduled for\s+(.+?)\.?\s*$", cleaned_text)
            if match:
                return match.group(1).strip()

        logger.debug(f"No value found for field '{field}'")
        return None

    def extract_multiline_value(
        self, paragraphs: List[Any], start_idx: int
    ) -> tuple[str, int]:
        """Extract multiline value starting from given index."""
        logger.debug(f"Extracting multiline value starting at index {start_idx}")
        content_lines = []
        i = start_idx + 1  # Start from next line

        # Special handling for TC number - only take first non-empty line
        if "TC number" in paragraphs[start_idx].text:
            while i < len(paragraphs):
                text = paragraphs[i].text.strip()
                if text:  # First non-empty line is our TC number
                    logger.debug(f"Found TC number: '{text}'")
                    return text, i
                i += 1
            return "", i - 1

        # Regular multiline field handling
        while i < len(paragraphs):
            text = paragraphs[i].text.strip()
            # Stop if we find a numbered header, known field, or blank line
            if (
                re.match(r"^\d+\..*", text)
                or any(
                    field in text for field in self.field_mappings if len(field) > 20
                )
                or not text
            ):  # Stop at blank line
                break

            logger.debug(f"Adding line {i}: '{text}'")
            content_lines.append(text)
            i += 1

        result = "\n".join(content_lines)
        logger.debug(f"Extracted multiline value: '{result}'")
        return result, i - 1

    def extract_conditional_block(
        self, paragraphs: List[Any], start_idx: int
    ) -> Optional[str]:
        """Extract selected option from conditional block."""
        logger.debug(f"Extracting conditional block starting at index {start_idx}")
        i = start_idx
        while i < len(paragraphs):
            text = paragraphs[i].text.strip()
            logger.debug(f"Checking conditional text: '{text}'")
            for option, value in self.checkbox_mappings.items():
                if "☒" in text and option.replace("☒", "").strip() in text:
                    logger.debug(f"Found checked option: '{value}'")
                    return value
            i += 1
        logger.debug("No checked option found in conditional block")
        return None

    def _extract_cfr_part(self, text: str) -> str:
        """Extract CFR part number from text."""
        logger.debug(f"Attempting to extract CFR part from: '{text}'")
        match = self.cfr_pattern.search(text)
        result = match.group(1) if match else ""
        logger.debug(f"Extracted CFR part: '{result}'")
        return result

    def extract_data(self, doc: Document) -> Dict[str, str]:
        """Extract all field values from the worksheet."""
        extracted_content = {
            placeholder: "" for _, placeholder in self.field_mappings.items()
        }
        extracted_content["{CFRPart}"] = ""

        logger.debug("Starting document extraction")

        paragraphs = list(doc.paragraphs)

        # Process CFR part first
        for paragraph in doc.paragraphs:
            if match := self.cfr_pattern.search(paragraph.text):
                extracted_content["{CFRPart}"] = match.group(1)
                break

        # Process multiline fields
        i = 0
        while i < len(doc.paragraphs):
            text = doc.paragraphs[i].text.strip()
            normalized = _normalize(text)

            # Handle multiline fields
            if any(field in normalized for field in self.multiline_fields):
                value, last_idx = self.extract_multiline_value(doc.paragraphs, i)
                for field, placeholder in self.field_mappings.items():
                    if field in normalized:
                        extracted_content[placeholder] = value
                i = last_idx
            i += 1

        i = 0
        while i < len(paragraphs):
            current_text = paragraphs[i].text.strip()
            logger.debug(f"Processing text: {current_text}")

            # Extract CFR Part
            if "CFR" in current_text or "Part" in current_text:
                extracted_content["{CFRPart}"] = self._extract_cfr_part(current_text)

            # Handle Q6 conditional block
            if "6. Check the appropriate box and complete:" in current_text:
                if project_type := self.extract_conditional_block(paragraphs, i + 1):
                    extracted_content["{ProjectType}"] = project_type
                i += 1
                continue

            # Handle multiline fields
            multiline_matched = False
            normalized_current = _normalize(current_text)
            for field in self.multiline_fields:
                if field in normalized_current:
                    value, last_idx = self.extract_multiline_value(paragraphs, i)
                    extracted_content[self.field_mappings[field]] = value
                    i = last_idx
                    multiline_matched = True
                    break

            if not multiline_matched:
                # Handle single-line fields with flexible matching
                current_normalized = _normalize(current_text)
                for field, placeholder in self.field_mappings.items():
                    if field in current_normalized:
                        value = (
                            current_normalized.split(field, 1)[1].lstrip(":").strip()
                        )
                        if value:
                            extracted_content[placeholder] = value
                            logger.debug(
                                "Extracted value: %s for %s", value, placeholder
                            )
                            break

            i += 1

        # Process tables
        for table in doc.tables:
            self.process_table(table, extracted_content)

        return extracted_content

    def process_table(self, table: Table, extracted_content: Dict[str, str]) -> None:
        """Process table to extract field values."""
        logger.debug("Processing table for field values")
        for row in table.rows:
            if len(row.cells) >= 2:
                cell_text = _normalize(self.clean_text(row.cells[0].text))
                cell_value = self.clean_text(row.cells[1].text)
                logger.debug(
                    f"Processing table cell - Label: '{cell_text}', Value: '{cell_value}'"
                )

                # Special handling for TC number - only take first line if multiline
                if "TC number" in cell_text:
                    cell_value = cell_value.split("\n")[0].strip()

                # Try exact match first
                for field, placeholder in self.field_mappings.items():
                    if field == cell_text:
                        extracted_content[placeholder] = cell_value
                        logger.info(
                            f"Found exact match {placeholder} in table: {cell_value}"
                        )
                        break

                # If no exact match, try contains match
                if not any(field == cell_text for field in self.field_mappings):
                    for field, placeholder in self.field_mappings.items():
                        if field in cell_text:
                            extracted_content[placeholder] = cell_value
                            logger.info(
                                f"Found partial match {placeholder} in table: {cell_value}"
                            )
                            break

        # Handle TC number specifically
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()

                tc_match = re.match(r".*TC number.*:\s*(.*)", text, re.IGNORECASE)
                if tc_match:
                    tc_value = tc_match.group(1).strip()
                    # If TC number is on same line
                    if tc_value:
                        extracted_content["{TCNumber}"] = tc_value
                    else:
                        # Look for TC number on next line
                        next_cell = self._get_next_cell(cell)
                        if next_cell:
                            next_text = next_cell.text.strip()
                            if next_text and not any(
                                key in next_text for key in FIELD_MAPPINGS
                            ):
                                extracted_content["{TCNumber}"] = next_text

    def _get_next_cell(self, cell: Any) -> Any:
        """Get next cell in document flow."""
        try:
            parent_row = cell._element.getparent()
            next_row = parent_row.getnext()
            if next_row is not None:
                return next_row.cells[0]
        except Exception:
            pass
        return None
