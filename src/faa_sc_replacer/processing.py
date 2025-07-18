"""Core processing logic for placeholder replacement."""

from __future__ import annotations

import re
from typing import Any, Dict

from docx import Document


def extract_fields(doc: Document) -> Dict[str, str]:
    """Extract placeholder values from a worksheet document."""

    field_mappings = {
        "Applicant name:": "{Applicant name}",
        "Airplane manufacturer:": "{Airplane manufacturer}",
        "Airplane model:": "{Airplane model}",
        "Subject of special conditions:": "{Subject of special conditions}",
        "Date of application:": "{Date of application}",
        "Type of airplane:": "{Type of airplane}",
        "TC number": "{TC number}",
        "Action prompting special conditions:": "{Action option}",
    }

    results: Dict[str, str] = {}

    def clean(text: str) -> str:
        return " ".join(text.strip().split())

    paragraphs = list(doc.paragraphs)
    for i, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        for field, placeholder in field_mappings.items():
            if text.startswith(field):
                value = text[len(field) :].strip()
                if not value and i + 1 < len(paragraphs):
                    value = paragraphs[i + 1].text.strip()
                results[placeholder] = clean(value)

    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                key = clean(row.cells[0].text)
                for field, placeholder in field_mappings.items():
                    if key.startswith(field):
                        value = clean(row.cells[1].text)
                        results[placeholder] = value

    return results


def replace_placeholders(doc: Document, values: Dict[str, str]) -> None:
    """Replace all placeholders in ``doc`` with provided values."""

    def process_paragraph(paragraph: Any) -> None:
        for run in paragraph.runs:
            for placeholder, val in values.items():
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, val)

    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

    for section in doc.sections:
        for hdr in (section.header, section.footer):
            for paragraph in hdr.paragraphs:
                process_paragraph(paragraph)
            for table in hdr.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            process_paragraph(paragraph)


def apply_conditionals(doc: Document, answers: Dict[str, str]) -> None:
    """Remove conditional blocks that do not match the provided option."""

    active = answers.get("{Action option}")
    if not active:
        return

    pattern = re.compile(r"\[\[OPTION_(\d)\]\](.*?)\[\[/OPTION_\1\]\]", re.DOTALL)

    def process_paragraph(paragraph: Any) -> None:
        for run in paragraph.runs:
            text = run.text
            if "[[OPTION_" in text:

                def repl(match: re.Match[str]) -> str:
                    option, content = match.group(1), match.group(2)
                    return content if option == active else ""

                run.text = re.sub(pattern, repl, text)

    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

    for section in doc.sections:
        for hdr in (section.header, section.footer):
            for paragraph in hdr.paragraphs:
                process_paragraph(paragraph)
            for table in hdr.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            process_paragraph(paragraph)
