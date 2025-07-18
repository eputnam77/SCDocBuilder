# Standard library imports
import os
import logging
from typing import Any, Dict, Optional
from datetime import datetime

# Third-party imports
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

# Set up logging
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)


class PlaceholderReplacer:
    """
    A class to handle replacement of placeholders in Word documents with content
    from Special Condition worksheets.
    """

    def __init__(self) -> None:
        """Initialize the PlaceholderReplacer with common placeholder mappings."""
        # Dictionary to store placeholder mappings
        self.placeholder_mappings: Dict[str, str] = {}

    def extract_worksheet_data(self, worksheet_path: str) -> Dict[str, str]:
        """
        Extract data from the Special Conditions Worksheet document.

        Args:
            worksheet_path (str): Path to the Special Conditions Worksheet document

        Returns:
            Dict[str, str]: Dictionary of field names and their values
        """
        try:
            # Load the worksheet document
            worksheet_doc = Document(worksheet_path)

            # Define fields to look for and their corresponding placeholders
            field_mappings = {
                "Applicant name:": "{Applicant name}",
                "Airplane manufacturer:": "{Airplane manufacturer}",
                "Airplane model:": "{Airplane model}",
                "Subject of special conditions:": "{Subject of special conditions}",
                "Date of application:": "{Date of application}",
                "Type of airplane:": "{Type of airplane}",
                "TC number": "{TC number}",
                "Name of SME:": "{Name of SME}",
                "Section name:": "{Section name}",
                "Routing symbol:": "{Routing symbol}",
                "SME Regional Office address:": "{SME Regional Office address}",
                "Telephone phone no:": "{Telephone phone no}",
                "E-mail:": "{E-mail}",
                "Briefly (one to three sentences) provide a summary of the novel or unusual design features of the airplane.": "{Summary}",
                "Provide a detailed discussion of the special conditions.": "{Description}",
            }

            extracted_content = {}

            # Function to clean extracted text
            def clean_text(text: str) -> str:
                """Remove extra whitespace and special characters."""
                return " ".join(text.strip().split())

            # Process paragraphs and tables to find field values
            for paragraph in worksheet_doc.paragraphs:
                text = paragraph.text.strip()
                for field, placeholder in field_mappings.items():
                    if text.startswith(field):
                        # Extract the value after the field label
                        value = text[len(field) :].strip()
                        if value:
                            extracted_content[placeholder] = clean_text(value)
                            logger.info(f"Found {placeholder}: {value}")

            # Process tables (some values might be in tables)
            for table in worksheet_doc.tables:
                for row in table.rows:
                    if len(row.cells) >= 2:  # Ensure we have at least 2 cells
                        cell_text = clean_text(row.cells[0].text)
                        for field, placeholder in field_mappings.items():
                            if cell_text.startswith(field):
                                value = clean_text(row.cells[1].text)
                                if value:
                                    extracted_content[placeholder] = value
                                    logger.info(
                                        f"Found {placeholder} in table: {value}"
                                    )

            # Log any missing fields
            for field, placeholder in field_mappings.items():
                if placeholder not in extracted_content:
                    logger.warning(f"Field '{field}' not found in worksheet")
                    extracted_content[placeholder] = ""

            return extracted_content

        except Exception as e:
            logger.error(f"Error extracting worksheet data: {str(e)}")
            raise

    def process_paragraph(
        self, paragraph: Paragraph, replacements: Dict[str, str]
    ) -> None:
        """
        Process a single paragraph for placeholder replacements.

        Args:
            paragraph (Paragraph): The paragraph to process
            replacements (Dict[str, str]): Dictionary of placeholder replacements
        """
        if not paragraph.text:
            return

        # Get the paragraph XML
        paragraph_xml = paragraph._element.xml

        # Make replacements in the XML
        for placeholder, replacement in replacements.items():
            if placeholder in paragraph_xml:
                # Preserve formatting by replacing within runs
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, replacement)

    def process_table(self, table: Table, replacements: Dict[str, str]) -> None:
        """
        Process a table for placeholder replacements.

        Args:
            table (Table): The table to process
            replacements (Dict[str, str]): Dictionary of placeholder replacements
        """
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self.process_paragraph(paragraph, replacements)

    def process_textboxes(self, part: Any, replacements: Dict[str, str]) -> None:
        """Process all text boxes within the given document part."""
        for txbx in part.element.xpath(".//w:txbxContent"):
            for p in txbx.xpath(".//w:p"):
                self.process_paragraph(Paragraph(p, part), replacements)

    def process_document(
        self, template_path: str, worksheet_path: str, output_path: Optional[str] = None
    ) -> None:
        """
        Process the entire Word document and replace placeholders.

        Args:
            template_path (str): Path to the Word document template
            worksheet_path (str): Path to the Special Conditions Worksheet
            output_path (Optional[str]): Path to save the processed document
        """
        try:
            # Validate input files exist
            if not os.path.exists(template_path):
                raise FileNotFoundError(f"Template file not found: {template_path}")
            if not os.path.exists(worksheet_path):
                raise FileNotFoundError(f"Worksheet file not found: {worksheet_path}")

            # Load the template document
            doc = Document(template_path)

            # Extract content from worksheet
            replacements = self.extract_worksheet_data(worksheet_path)

            # Process all paragraphs, tables, and text boxes including headers/footers
            for paragraph in doc.paragraphs:
                self.process_paragraph(paragraph, replacements)
            self.process_textboxes(doc, replacements)

            for table in doc.tables:
                self.process_table(table, replacements)

            for section in doc.sections:
                for hdr in (section.header, section.footer):
                    for paragraph in hdr.paragraphs:
                        self.process_paragraph(paragraph, replacements)
                    self.process_textboxes(hdr, replacements)
                    for table in hdr.tables:
                        self.process_table(table, replacements)

            # Generate output path if not provided
            if output_path is None:
                file_name = os.path.basename(template_path)
                base_name, ext = os.path.splitext(file_name)
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_path = f"{base_name}_processed_{timestamp}{ext}"

            # Save the processed document
            doc.save(output_path)
            logger.info(f"Document processed and saved to: {output_path}")

        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            raise


def main() -> None:
    """Main function to demonstrate usage of the PlaceholderReplacer class."""
    try:
        # Initialize the replacer
        replacer = PlaceholderReplacer()

        # Set paths to your actual files
        template_path = r"C:\Users\Hoctar\Downloads\SC - Notice of Proposed SC TEMPLATE - Placeholders.docx"  # The document with placeholders
        worksheet_path = r"C:\Users\Hoctar\Downloads\SC worksheet_Airbus_AT11885IB-T_49degreeobliqueseats_sml101724.docx"  # The Special Conditions Worksheet
        output_path = r"C:\Users\Hoctar\Downloads\SC worksheet_Airbus_AT11885IB-T_49degreeobliqueseats_sml101724-processed.docx"  # Where to save the result (optional)

        # Process the document
        replacer.process_document(
            template_path=template_path,
            worksheet_path=worksheet_path,
            output_path=output_path,
        )

    except Exception as e:
        logger.error(f"Error in main function: {str(e)}")
        raise
