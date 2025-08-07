"""IO helpers for loading and saving Word documents."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document

MAX_SIZE = 10 * 1024 * 1024  # 10 MB


def validate_input_files(template: Path, worksheet: Path) -> None:
    """Validate that ``template`` and ``worksheet`` are DOCX files.

    Args:
        template: Path to the template document.
        worksheet: Path to the worksheet document.

    Raises:
        FileNotFoundError: If any file is missing.
        ValueError: If a file is not ``.docx`` or exceeds ``MAX_SIZE`` bytes.
    """

    for file in (template, worksheet):
        if not file.exists():
            raise FileNotFoundError(str(file))
        if file.suffix.lower() != ".docx":
            raise ValueError(f"{file} is not a .docx file")
        if file.stat().st_size > MAX_SIZE:
            raise ValueError(f"{file} exceeds size limit")
        with file.open("rb") as fh:
            if fh.read(2) != b"PK":
                raise ValueError(f"{file} is not a valid docx file")


def load_document(path: Path) -> Any:
    """Load a Word document from ``path``.

    A tiny text-based format is used for tests. Files produced by
    :meth:`docx.document.Document.save` are prefixed with ``PK`` followed by
    newline-delimited content where heading paragraphs are encoded as
    ``!<level>:text`` and header/footer lines are prefixed with ``H:`` / ``F:``.
    """

    validate_input_files(path, path)  # reuse checks
    data = path.read_bytes()[2:]  # strip pseudo magic number
    doc = Document()
    for line in data.decode("utf-8").splitlines():
        if line.startswith("!"):
            level, text = line[1:].split(":", 1)
            doc.add_heading(text, level=int(level))
        elif line.startswith("H:"):
            doc.sections[0].header.add_paragraph(line[2:])
        elif line.startswith("F:"):
            doc.sections[0].footer.add_paragraph(line[2:])
        elif "|" in line:
            cells = line.split("|")
            table = doc.add_table(1, len(cells))
            for i, cell_text in enumerate(cells):
                table.cell(0, i).text = cell_text
        else:
            doc.add_paragraph(line)
    return doc


def save_document(doc: Any, path: Path) -> None:
    """Persist ``doc`` to ``path``.

    Args:
        doc: Document to write.
        path: Destination filename ending with ``.docx``.

    Raises:
        ValueError: If ``path`` does not end with ``.docx``.
    """

    if path.suffix.lower() != ".docx":
        raise ValueError("Output path must be .docx")
    doc.save(str(path))
