"""Core processing logic for placeholder replacement."""

from __future__ import annotations

import re
from docx.document import Document
from docx.text.paragraph import Paragraph

from typing import Any, Dict

DEFAULT_FIELD_MAPPINGS: Dict[str, str] = {
    "Applicant name:": "{Applicant name}",
    "Airplane manufacturer:": "{Airplane manufacturer}",
    "Airplane model:": "{Airplane model}",
    "Subject of special conditions:": "{Subject of special conditions}",
    "Date of application:": "{Date of application}",
    "Type of airplane:": "{Type of airplane}",
    "TC number": "{TC number}",
    "Action prompting special conditions:": "{Action option}",
}

# Precompile regex used for conditional blocks once at module import
OPTION_PATTERN = re.compile(r"\[\[OPTION_(\d)\]\](.*?)\[\[/OPTION_\1\]\]", re.DOTALL)


def _iter_textbox_paragraphs(part: Any) -> list[Paragraph]:
    """Return paragraphs contained in text boxes of ``part``."""
    paragraphs: list[Paragraph] = []
    for txbx in part.element.xpath(".//w:txbxContent"):
        for p in txbx.xpath(".//w:p"):
            paragraphs.append(Paragraph(p, part))
    return paragraphs


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """Replace paragraph runs with a single run containing ``text``."""

    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    if text:
        paragraph.add_run(text)


def extract_fields(
    doc: Document, field_mappings: Dict[str, str] | None = None
) -> Dict[str, str]:
    """Extract placeholder values from a worksheet document.

    ``field_mappings`` allows loading custom schemas from configuration files.
    """

    if field_mappings is None:
        field_mappings = DEFAULT_FIELD_MAPPINGS

    results: Dict[str, str] = {}

    def clean(text: str) -> str:
        return text.strip()

    paragraphs = list(doc.paragraphs)
    for i, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        for field, placeholder in field_mappings.items():
            if text.startswith(field):
                value = text[len(field) :].strip()
                if not value:
                    lines = []
                    j = i + 1
                    while j < len(paragraphs):
                        next_text = paragraphs[j].text.strip()
                        if not next_text or any(
                            next_text.startswith(f) for f in field_mappings
                        ):
                            break
                        lines.append(next_text)
                        j += 1
                    value = "\n".join(lines)
                results[placeholder] = clean(value)

    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                key = clean(row.cells[0].text)
                for field, placeholder in field_mappings.items():
                    if key.startswith(field):
                        value = clean(row.cells[1].text)
                        results[placeholder] = value

    return results


def replace_placeholders(doc: Document, values: Dict[str, str]) -> None:
    """Replace all placeholders in ``doc`` with provided values."""

    def process_paragraph(paragraph: Any) -> None:
        for run in paragraph.runs:
            for placeholder, val in values.items():
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, val)

    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)
    for paragraph in _iter_textbox_paragraphs(doc.part):
        process_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

    for section in doc.sections:
        for hdr in (section.header, section.footer):
            for paragraph in hdr.paragraphs:
                process_paragraph(paragraph)
            for table in hdr.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            process_paragraph(paragraph)
            for paragraph in _iter_textbox_paragraphs(hdr.part):
                process_paragraph(paragraph)


def apply_conditionals(doc: Document, answers: Dict[str, str]) -> None:
    """Remove conditional blocks that do not match the provided option."""

    active = answers.get("{Action option}")
    if not active:
        return

    def process_paragraph(paragraph: Paragraph) -> None:
        full_text = "".join(run.text for run in paragraph.runs)
        if "[[OPTION_" not in full_text:
            return

        def repl(match: re.Match[str]) -> str:
            option, content = match.group(1), match.group(2)
            return content if option == active else ""

        new_text = re.sub(OPTION_PATTERN, repl, full_text)
        if new_text != full_text:
            cleaned = "\n".join(
                line.strip() for line in new_text.splitlines() if line.strip()
            )
            _set_paragraph_text(paragraph, cleaned)

    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)
    for paragraph in _iter_textbox_paragraphs(doc.part):
        process_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

    for section in doc.sections:
        for hdr in (section.header, section.footer):
            for paragraph in hdr.paragraphs:
                process_paragraph(paragraph)
            for table in hdr.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            process_paragraph(paragraph)
            for paragraph in _iter_textbox_paragraphs(hdr.part):
                process_paragraph(paragraph)
