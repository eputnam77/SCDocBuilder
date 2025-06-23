from pathlib import Path
from docx import Document

from faa_sc_filler.streamlit_app import process_files


def create_docs(tmp_path):
    template_path = tmp_path / "template.docx"
    worksheet_path = tmp_path / "worksheet.docx"
    template = Document()
    template.add_paragraph("Contact: {SMEName}")
    template.save(template_path)

    worksheet = Document()
    worksheet.add_paragraph("Name of SME: John Doe")
    worksheet.save(worksheet_path)

    return template_path.read_bytes(), worksheet_path.read_bytes()


def test_process_files(tmp_path):
    template_bytes, worksheet_bytes = create_docs(tmp_path)
    output = process_files(template_bytes, worksheet_bytes)
    assert output is not None
    assert Path(output).exists()
