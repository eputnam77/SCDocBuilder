from __future__ import annotations

from datetime import datetime, tzinfo
from pathlib import Path
from typing import Any, cast
import logging
from logging.handlers import RotatingFileHandler
import json
import runpy
import sys
import scdocbuilder.cli

import pytest
import typing

if typing.TYPE_CHECKING:
    from docx import Document
else:
    pytest.importorskip("docx")
    from docx import Document

from scdocbuilder.cli import ErrorCode, main, parse_args


def test_parse_args_parses_required() -> None:
    args = parse_args(
        [
            "--template",
            "t.docx",
            "--worksheet",
            "w.docx",
            "--dry-run",
            "--log-level",
            "DEBUG",
        ]
    )
    assert args.template == "t.docx"
    assert args.worksheet == "w.docx"
    assert args.dry_run is True
    assert args.log_level == "DEBUG"


def test_parse_args_requires_template() -> None:
    with pytest.raises(SystemExit):
        parse_args(["--worksheet", "w.docx"])


def test_parse_args_requires_input() -> None:
    with pytest.raises(SystemExit):
        parse_args(["--template", "t.docx"])


def test_parse_args_allows_show_completion_without_required() -> None:
    args = parse_args(["--show-completion"])
    assert args.show_completion == "bash"


def test_generate_completion_invalid_shell() -> None:
    with pytest.raises(ValueError):
        scdocbuilder.cli._generate_completion("fish")


def test_main_show_completion_invalid_shell() -> None:
    with pytest.raises(SystemExit) as exc:
        main(["--show-completion", "fish"])
    assert exc.value.code == ErrorCode.EVALID


def test_main_show_completion_prints(capsys: Any) -> None:
    main(["--show-completion"])
    out, err = capsys.readouterr()
    assert "scdocbuilder" in out


def test_main_batch_missing_directory(tmp_path: Path) -> None:
    template = tmp_path / "t.docx"
    Document().save(str(template))
    missing = tmp_path / "missing"
    with pytest.raises(SystemExit) as exc:
        main(["--template", str(template), "--batch", str(missing)])
    assert exc.value.code == ErrorCode.ENOFILE


def test_cli_module_entrypoint(monkeypatch: pytest.MonkeyPatch, capsys: Any) -> None:
    monkeypatch.setattr(sys, "argv", ["scdocbuilder.cli", "--show-completion"])
    runpy.run_module("scdocbuilder.cli", run_name="__main__")
    out, _ = capsys.readouterr()
    assert "scdocbuilder" in out


def test_parse_args_parses_batch() -> None:
    args = parse_args(
        [
            "--template",
            "t.docx",
            "--batch",
            "worksheets",
        ]
    )

    assert args.template == "t.docx"
    assert args.batch == "worksheets"
    assert args.worksheet is None


def test_cli_exits_zero_with_required_args(tmp_path: Path, capsys: Any) -> None:
    """Running the CLI with required arguments should produce an output file."""

    template = tmp_path / "t.docx"
    worksheet = tmp_path / "w.docx"
    Document().save(str(template))
    ws_doc = Document()
    ws_doc.add_paragraph("Applicant name: Foo")
    ws_doc.add_paragraph("Airplane model: Bar")
    ws_doc.add_paragraph("Question 15:")
    ws_doc.add_paragraph("Ans15")
    ws_doc.add_paragraph("Question 16:")
    ws_doc.add_paragraph("Ans16")
    ws_doc.add_paragraph("Question 17:")
    ws_doc.add_paragraph("Ans17")
    ws_doc.save(str(worksheet))

    main(
        [
            "--template",
            str(template),
            "--worksheet",
            str(worksheet),
        ]
    )
    output_path = Path(capsys.readouterr().out.strip())
    assert output_path.exists()


def test_main_prints_output_path(tmp_path: Path, capsys: Any, monkeypatch: Any) -> None:
    """CLI should print the absolute output path when no --output is given."""

    template = tmp_path / "t.docx"
    worksheet = tmp_path / "w.docx"
    Document().save(str(template))
    ws_doc = Document()
    ws_doc.add_paragraph("Applicant name: Foo")
    ws_doc.add_paragraph("Airplane model: Bar")
    ws_doc.add_paragraph("Question 15:")
    ws_doc.add_paragraph("Ans15")
    ws_doc.add_paragraph("Question 16:")
    ws_doc.add_paragraph("Ans16")
    ws_doc.add_paragraph("Question 17:")
    ws_doc.add_paragraph("Ans17")
    ws_doc.save(str(worksheet))

    fixed_time = datetime(2024, 1, 2, 3, 4, 5)

    class FixedDateTime(datetime):
        @classmethod
        def now(cls, tz: tzinfo | None = None) -> FixedDateTime:
            return cast(FixedDateTime, fixed_time)

    monkeypatch.chdir(tmp_path)
    monkeypatch.setattr("scdocbuilder.cli.datetime", FixedDateTime)

    main(
        [
            "--template",
            str(template),
            "--worksheet",
            str(worksheet),
        ]
    )

    expected = tmp_path / f"t_{fixed_time:%Y%m%d_%H%M%S}.docx"
    assert expected.exists()
    captured = capsys.readouterr()
    assert captured.out.strip() == str(expected.resolve())


def test_main_respects_output_flag(tmp_path: Path, capsys: Any) -> None:
    """CLI should write to the path provided by --output."""

    template = tmp_path / "t.docx"
    worksheet = tmp_path / "w.docx"
    Document().save(str(template))
    ws_doc = Document()
    ws_doc.add_paragraph("Applicant name: Foo")
    ws_doc.add_paragraph("Airplane model: Bar")
    ws_doc.add_paragraph("Question 15:")
    ws_doc.add_paragraph("Ans15")
    ws_doc.add_paragraph("Question 16:")
    ws_doc.add_paragraph("Ans16")
    ws_doc.add_paragraph("Question 17:")
    ws_doc.add_paragraph("Ans17")
    ws_doc.save(str(worksheet))
    output = tmp_path / "custom.docx"

    main(
        [
            "--template",
            str(template),
            "--worksheet",
            str(worksheet),
            "--output",
            str(output),
        ]
    )

    assert output.exists()
    captured = capsys.readouterr()
    assert captured.out.strip() == str(output.resolve())


def test_main_missing_file_exits_with_code(tmp_path: Path) -> None:
    """CLI should exit with ENOFILE when files are missing."""

    missing = tmp_path / "missing.docx"
    with pytest.raises(SystemExit) as exc:
        main(
            [
                "--template",
                str(missing),
                "--worksheet",
                str(missing),
            ]
        )

    assert exc.value.code == ErrorCode.ENOFILE


def test_main_batch_processes_directory(tmp_path: Path, monkeypatch: Any) -> None:
    template = tmp_path / "t.docx"
    Document().save(str(template))
    batch_dir = tmp_path / "ws"
    batch_dir.mkdir()
    for i in range(2):
        doc = Document()
        doc.add_paragraph("Applicant name: Foo")
        doc.add_paragraph("Airplane model: Bar")
        doc.add_paragraph("Question 15:")
        doc.add_paragraph("Ans15")
        doc.add_paragraph("Question 16:")
        doc.add_paragraph("Ans16")
        doc.add_paragraph("Question 17:")
        doc.add_paragraph("Ans17")
        doc.save(str(batch_dir / f"w{i}.docx"))

    fixed_time = datetime(2024, 1, 1, 1, 1, 1)

    class FixedDateTime(datetime):
        @classmethod
        def now(cls, tz: tzinfo | None = None) -> FixedDateTime:
            return cast(FixedDateTime, fixed_time)

    monkeypatch.setattr("scdocbuilder.cli.datetime", FixedDateTime)
    monkeypatch.chdir(tmp_path)

    main(
        [
            "--template",
            str(template),
            "--batch",
            str(batch_dir),
        ]
    )

    for i in range(2):
        expected = batch_dir / f"w{i}_{fixed_time:%Y%m%d_%H%M%S}.docx"
        assert expected.exists()


def test_logging_rotation_configured(tmp_path: Path, monkeypatch: Any) -> None:
    template = tmp_path / "t.docx"
    worksheet = tmp_path / "w.docx"
    Document().save(str(template))
    ws_doc = Document()
    ws_doc.add_paragraph("Applicant name: Foo")
    ws_doc.add_paragraph("Airplane model: Bar")
    ws_doc.add_paragraph("Question 15:")
    ws_doc.add_paragraph("Ans15")
    ws_doc.add_paragraph("Question 16:")
    ws_doc.add_paragraph("Ans16")
    ws_doc.add_paragraph("Question 17:")
    ws_doc.add_paragraph("Ans17")
    ws_doc.save(str(worksheet))

    monkeypatch.chdir(tmp_path)
    logging.getLogger().handlers.clear()
    main(
        [
            "--template",
            str(template),
            "--worksheet",
            str(worksheet),
        ]
    )

    handlers = logging.getLogger().handlers
    assert any(
        isinstance(h, RotatingFileHandler) and h.maxBytes == 5 * 1024 * 1024
        for h in handlers
    )
    logging.getLogger().handlers.clear()


def test_main_dry_run(tmp_path: Path, capsys: Any) -> None:
    template = tmp_path / "t.docx"
    worksheet = tmp_path / "w.docx"
    doc = Document()
    doc.add_paragraph("{Applicant name}")
    doc.save(str(template))
    ws = Document()
    ws.add_paragraph("Applicant name: Foo")
    ws.add_paragraph("Airplane model: Bar")
    ws.add_paragraph("Question 15:")
    ws.add_paragraph("Ans15")
    ws.add_paragraph("Question 16:")
    ws.add_paragraph("Ans16")
    ws.add_paragraph("Question 17:")
    ws.add_paragraph("Ans17")
    ws.save(str(worksheet))

    main(
        [
            "--template",
            str(template),
            "--worksheet",
            str(worksheet),
            "--dry-run",
        ]
    )

    out = capsys.readouterr().out
    assert '"new": "Foo"' in out


def test_main_batch_dry_run(tmp_path: Path, capsys: Any) -> None:
    template = tmp_path / "t.docx"
    Document().save(str(template))
    batch = tmp_path / "b"
    batch.mkdir()
    for i in range(2):
        ws = Document()
        ws.add_paragraph("Applicant name: Foo")
        ws.add_paragraph("Airplane model: Bar")
        ws.add_paragraph("Question 15:")
        ws.add_paragraph("Ans15")
        ws.add_paragraph("Question 16:")
        ws.add_paragraph("Ans16")
        ws.add_paragraph("Question 17:")
        ws.add_paragraph("Ans17")
        ws.save(str(batch / f"w{i}.docx"))

    main(
        [
            "--template",
            str(template),
            "--batch",
            str(batch),
            "--dry-run",
        ]
    )

    out = capsys.readouterr().out
    assert '"new": "Foo"' in out


def test_main_with_schema(tmp_path: Path, capsys: Any) -> None:
    template = tmp_path / "t.docx"
    doc = Document()
    doc.add_paragraph("{Applicant name} {X}")
    doc.save(str(template))
    worksheet = tmp_path / "w.docx"
    ws = Document()
    ws.add_paragraph("Applicant name: Foo")
    ws.add_paragraph("Airplane model: Baz")
    ws.add_paragraph("X: Bar")
    ws.add_paragraph("Question 15:")
    ws.add_paragraph("Ans15")
    ws.add_paragraph("Question 16:")
    ws.add_paragraph("Ans16")
    ws.add_paragraph("Question 17:")
    ws.add_paragraph("Ans17")
    ws.save(str(worksheet))
    schema = tmp_path / "schema.json"
    schema.write_text('{"X:": "{X}"}')

    main(
        [
            "--template",
            str(template),
            "--worksheet",
            str(worksheet),
            "--schema",
            str(schema),
        ]
    )

    out_path = Path(capsys.readouterr().out.strip())
    assert out_path.exists()
    processed = Document(str(out_path))
    assert "Bar" in processed.paragraphs[0].text


def test_main_html_out(tmp_path: Path) -> None:
    template = tmp_path / "t.docx"
    doc = Document()
    doc.add_paragraph("{Applicant name}")
    doc.save(str(template))
    worksheet = tmp_path / "w.docx"
    ws = Document()
    ws.add_paragraph("Applicant name: Foo")
    ws.add_paragraph("Airplane model: Bar")
    ws.add_paragraph("Question 15:")
    ws.add_paragraph("Ans15")
    ws.add_paragraph("Question 16:")
    ws.add_paragraph("Ans16")
    ws.add_paragraph("Question 17:")
    ws.add_paragraph("Ans17")
    ws.save(str(worksheet))
    html_out = tmp_path / "out.html"

    main(
        [
            "--template",
            str(template),
            "--worksheet",
            str(worksheet),
            "--html-out",
            str(html_out),
        ]
    )

    assert html_out.exists()
    assert "<p" in html_out.read_text()


def test_main_dry_run_outputs_json_and_no_file(
    tmp_path: Path, capsys: Any, monkeypatch: Any
) -> None:
    template = tmp_path / "t.docx"
    doc = Document()
    doc.add_paragraph("{Applicant name}")
    doc.save(str(template))
    worksheet = tmp_path / "w.docx"
    ws = Document()
    ws.add_paragraph("Applicant name: Foo")
    ws.add_paragraph("Airplane model: Bar")
    ws.add_paragraph("Question 15:")
    ws.add_paragraph("Ans15")
    ws.add_paragraph("Question 16:")
    ws.add_paragraph("Ans16")
    ws.add_paragraph("Question 17:")
    ws.add_paragraph("Ans17")
    ws.save(str(worksheet))
    fixed_time = datetime(2024, 2, 3, 4, 5, 6)

    class FixedDateTime(datetime):
        @classmethod
        def now(cls, tz: tzinfo | None = None) -> FixedDateTime:
            return cast(FixedDateTime, fixed_time)

    monkeypatch.setattr("scdocbuilder.cli.datetime", FixedDateTime)
    monkeypatch.chdir(tmp_path)

    main(
        [
            "--template",
            str(template),
            "--worksheet",
            str(worksheet),
            "--dry-run",
        ]
    )

    out = capsys.readouterr().out
    diff = json.loads(out)
    assert diff["{Applicant name}"]["old"] == "{Applicant name}"
    assert diff["{Applicant name}"]["new"] == "Foo"
    expected = tmp_path / f"t_{fixed_time:%Y%m%d_%H%M%S}.docx"
    assert not expected.exists()


def test_main_validation_error_exit_code(tmp_path: Path) -> None:
    template = tmp_path / "t.docx"
    doc = Document()
    doc.add_paragraph("{Applicant name}")
    doc.save(str(template))
    worksheet = tmp_path / "w.docx"
    ws = Document()
    ws.add_paragraph("Applicant name: Foo")
    ws.save(str(worksheet))

    with pytest.raises(SystemExit) as exc:
        main(
            [
                "--template",
                str(template),
                "--worksheet",
                str(worksheet),
            ]
        )
    assert exc.value.code == ErrorCode.EVALID


def test_main_wrong_mime_exit_code(tmp_path: Path) -> None:
    template = tmp_path / "t.docx"
    Document().save(str(template))
    worksheet = tmp_path / "w.docx"
    worksheet.write_text("not a real docx")

    with pytest.raises(SystemExit) as exc:
        main(
            [
                "--template",
                str(template),
                "--worksheet",
                str(worksheet),
            ]
        )
    assert exc.value.code == ErrorCode.EVALID


def test_main_replacement_failure_exit_code(tmp_path: Path, monkeypatch: Any) -> None:
    template = tmp_path / "t.docx"
    worksheet = tmp_path / "w.docx"
    doc = Document()
    doc.add_paragraph("{Applicant name}")
    doc.save(str(template))
    ws = Document()
    ws.add_paragraph("Applicant name: Foo")
    ws.add_paragraph("Airplane model: Bar")
    ws.add_paragraph("Question 15:")
    ws.add_paragraph("Ans15")
    ws.add_paragraph("Question 16:")
    ws.add_paragraph("Ans16")
    ws.add_paragraph("Question 17:")
    ws.add_paragraph("Ans17")
    ws.save(str(worksheet))

    def boom(*args: Any, **kwargs: Any) -> None:
        raise RuntimeError("boom")

    monkeypatch.setattr("scdocbuilder.processing.replace_placeholders", boom)

    with pytest.raises(SystemExit) as exc:
        main(
            [
                "--template",
                str(template),
                "--worksheet",
                str(worksheet),
            ]
        )
    assert exc.value.code == ErrorCode.EREPLACE


def test_cli_shows_completion_script() -> None:
    """CLI should expose completion script via --show-completion."""
    script = scdocbuilder.cli._generate_completion("bash")
    assert "--template" in script
    assert "--worksheet" in script
