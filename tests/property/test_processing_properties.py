from pathlib import Path

import typing
import pytest

if not typing.TYPE_CHECKING:
    pytest.importorskip("docx")

from scdocbuilder.processing import (
    extract_fields,
    replace_placeholders,
    apply_conditionals,
)

pytest.importorskip("hypothesis")


@pytest.mark.property
def test_extract_fields_roundtrip(tmp_path: Path) -> None:
    """extract_fields should return mappings from worksheet."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Applicant name: Test Co")
    doc.add_paragraph("Action prompting special conditions:")
    doc.add_paragraph("1")
    path = tmp_path / "w.docx"
    doc.save(str(path))

    loaded = Document(str(path))
    fields = extract_fields(loaded)
    assert fields["{Applicant name}"] == "Test Co"
    assert fields["{Action option}"] == "1"


@pytest.mark.property
def test_replace_placeholders_updates_doc(tmp_path: Path) -> None:
    """replace_placeholders should modify document text."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Hello {name}")
    values = {"{name}": "World"}
    replace_placeholders(doc, values)
    assert "World" in doc.paragraphs[0].text


@pytest.mark.property
def test_apply_conditionals_keeps_selected(tmp_path: Path) -> None:
    """Only the selected OPTION block should remain."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("[[OPTION_1]]A[[/OPTION_1]][[OPTION_2]]B[[/OPTION_2]]")
    apply_conditionals(doc, {"{Action option}": "1"})
    assert doc.paragraphs[0].text == "A"
