# pytest tests/test_replacer.py -v

import pytest
from docx import Document
import logging
from faa_sc_filler.replacer import PlaceholderReplacer

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def test_placeholder_replacement():
    # Create test document with known placeholder
    doc = Document()
    doc.add_paragraph("Test SME Name: {SMEName}")
    
    replacer = PlaceholderReplacer()
    replacements = {"{SMEName}": "John Doe"}
    
    # Process the paragraph
    replacer.process_paragraph(doc.paragraphs[0], replacements)
    
    # Verify replacement
    assert "John Doe" in doc.paragraphs[0].text
    assert "{SMEName}" not in doc.paragraphs[0].text

def test_split_placeholder():
    """Test replacing a placeholder split across multiple runs."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Contact: {")
    p.add_run("SME")
    p.add_run("Name}")
    
    replacer = PlaceholderReplacer()
    replacements = {"{SMEName}": "John Doe"}
    
    replacer.process_paragraph(p, replacements)
    assert "John Doe" in p.text
    assert "{SMEName}" not in p.text

def test_multiple_replacements_in_run():
    """Test multiple replacements in a single run."""
    doc = Document()
    p = doc.add_paragraph("Name: {SMEName}, Email: {SMEEmail}")
    
    replacer = PlaceholderReplacer()
    replacements = {
        "{SMEName}": "John Doe",
        "{SMEEmail}": "john@example.com"
    }
    
    replacer.process_paragraph(p, replacements)
    assert "John Doe" in p.text
    assert "john@example.com" in p.text
    assert "{SMEName}" not in p.text
    assert "{SMEEmail}" not in p.text

def test_preserve_formatting():
    """Test that formatting is preserved after replacement."""
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("{SMEName}")
    run.bold = True
    run.italic = True
    
    replacer = PlaceholderReplacer()
    replacements = {"{SMEName}": "John Doe"}
    
    replacer.process_paragraph(p, replacements)
    assert "John Doe" in p.text
    assert p.runs[0].bold
    assert p.runs[0].italic

def test_empty_runs():
    """Test handling of empty runs in paragraph."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("")  # Empty run
    p.add_run("{SMEName}")
    p.add_run("")  # Empty run
    
    replacer = PlaceholderReplacer()
    replacements = {"{SMEName}": "John Doe"}
    
    replacer.process_paragraph(p, replacements)
    assert "John Doe" in p.text

def test_partial_placeholder():
    """Test handling of partial/malformed placeholders."""
    doc = Document()
    p = doc.add_paragraph("Contact: {SMEName, Phone: SMEPhone}")  # Missing closing brace
    
    replacer = PlaceholderReplacer()
    replacements = {"{SMEName}": "John Doe"}
    
    replacer.process_paragraph(p, replacements)
    assert p.text == "Contact: {SMEName, Phone: SMEPhone}"  # Should not modify malformed placeholders

def test_comprehensive_replacement():
    """Test multiple types of replacements in a single document."""
    doc = Document()
    p1 = doc.add_paragraph("SME: {SMEName}, Phone: {SMEPhone}")
    p2 = doc.add_paragraph()
    p2.add_run("Email: {").bold = True
    p2.add_run("SMEEmail").italic = True
    p2.add_run("}")
    
    replacer = PlaceholderReplacer()
    replacements = {
        "{SMEName}": "John Doe",
        "{SMEPhone}": "123-456-7890",
        "{SMEEmail}": "john@example.com"
    }
    
    # Process paragraphs
    for p in doc.paragraphs:
        replacer.process_paragraph(p, replacements)
        logger.debug(f"Processed paragraph: {p.text}")
    
    # Verify results
    assert "John Doe" in p1.text
    assert "123-456-7890" in p1.text
    assert "john@example.com" in p2.text
    assert p2.runs[0].bold  # Formatting preserved
    assert p2.runs[1].italic  # Formatting preserved

def test_multi_paragraph_replacement():
    """Test replacing a placeholder with multiple paragraphs of text."""
    doc = Document()
    template_text = "The Proposed Special Conditions\nAccordingly, the Federal Aviation Administration (FAA) proposes the following special conditions as part of the type certification basis for {AirplaneManufacturer} Model {AirplaneModel} if applicable: series airplanes, as modified by {ApplicantName}.\n\n{SpecialConditions}"
    doc.add_paragraph(template_text)
    
    replacer = PlaceholderReplacer()
    special_conditions = """Section 25.785(d) requires that each occupant of a seat installed at an angle of more than 18 degrees, relative to bow-to-stern airplane cabin centerline, must be protected from head injury using a seatbelt and an energy-absorbing rest that supports the arms, shoulders, head, and spine, or using a seatbelt and shoulder harness designed to prevent the head from contacting any injurious object.

The Airbus Model A321 neo ACF and A321 neo XLR airplane's single occupant oblique seat installation with airbag devices and 3-point restraint or pretensioner restraint system is novel such that the current requirements do not adequately address airbag or pretensioner devices and protection of the occupant's neck, spine torso, and legs for seating configurations that are positioned at an angle of 49 degrees from the airplane centerline. The seating configuration installation angle is beyond the installation-design limits of current special conditions issued for seat positions at angles between 18 degrees and 45 degrees. For example, at these angles, lateral neck bending and other injury mechanisms prevalent from a fully side-facing installation become a concern. To account for these concerns, these special conditions are based on FAA policy statement PS-AIR-25-27, "Technical Criteria for Approving Obliques seats" as well as policy statement PS-ANM-25-03-R1, "Technical Criteria for Approving Side-Facing Seats.\""""
    
    replacements = {
        "{AirplaneManufacturer}": "Airbus",
        "{AirplaneModel}": "A321 neo ACF and A321 neo XLR",
        "{ApplicantName}": "STC Holder",
        "{SpecialConditions}": special_conditions
    }
    
    replacer.process_paragraph(doc.paragraphs[0], replacements)
    
    # Verify all replacements were made
    result_text = doc.paragraphs[0].text
    assert "{SpecialConditions}" not in result_text
    assert "{AirplaneManufacturer}" not in result_text
    assert "{AirplaneModel}" not in result_text
    assert "{ApplicantName}" not in result_text
    assert "Section 25.785(d)" in result_text
    assert "Airbus Model A321 neo ACF and A321 neo XLR" in result_text
    assert "PS-AIR-25-27" in result_text
