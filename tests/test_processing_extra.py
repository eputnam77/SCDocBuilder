from typing import Any
from types import MethodType
import typing
import pytest
import re
from pathlib import Path

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

if typing.TYPE_CHECKING:
    from docx import Document
    from docx.shared import Inches
else:
    pytest.importorskip("docx")
    from docx import Document
    from docx.shared import Inches

from scdocbuilder import processing
from scdocbuilder.processing import (
    _iter_textbox_paragraphs,
    _set_paragraph_text,
    replace_placeholders,
    apply_conditionals,
)


def _patch_iter_textbox(monkeypatch: Any, text: str) -> None:
    para = Document().add_paragraph(text)
    monkeypatch.setattr(processing, "_iter_textbox_paragraphs", lambda part: [para])


def _patch_textbox(monkeypatch: Any, doc: Any, text: str) -> None:
    # Only used for testing _iter_textbox_paragraphs itself
    para = Document().add_paragraph(text)
    p_el = para._p

    class FakeTbx:
        def xpath(self, query: str) -> list[Any]:
            return [p_el]

    def fake_xpath(self: Any, query: str) -> list[FakeTbx]:
        return [FakeTbx()]

    monkeypatch.setattr(
        doc.part.element,
        "xpath",
        MethodType(fake_xpath, doc.part.element),
    )


def test_iter_textbox_paragraphs(monkeypatch: Any) -> None:
    doc = Document()
    _patch_textbox(monkeypatch, doc, "Box")
    paras = _iter_textbox_paragraphs(doc.part)
    assert [p.text for p in paras] == ["Box"]


def test_set_paragraph_text() -> None:
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("A")
    p.add_run("B")
    _set_paragraph_text(p, "C")
    assert len(p.runs) == 1
    assert p.text == "C"


def test_extract_fields_table() -> None:
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Applicant name:"
    table.cell(0, 1).text = "Foo"
    fields = processing.extract_fields(doc)
    assert fields["{Applicant name}"] == "Foo"


def test_extract_fields_multiline_answer() -> None:
    doc = Document()
    doc.add_paragraph("Question 15:")
    doc.add_paragraph("Line1")
    doc.add_paragraph("Line2")
    fields = processing.extract_fields(doc, {"Question 15:": "{Q15}"})
    assert fields["{Q15}"] == "Line1\nLine2"


def test_replace_placeholders_full(monkeypatch: Any) -> None:
    doc = Document()
    doc.add_paragraph("{x}")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "{x}"
    doc.sections[0].header.add_paragraph("{x}")
    hdr_table = doc.sections[0].header.add_table(rows=1, cols=1, width=Inches(1))
    hdr_table.cell(0, 0).text = "{x}"
    _patch_iter_textbox(monkeypatch, "{x}")

    replace_placeholders(doc, {"{x}": "VAL"})

    texts = [para.text for para in doc.paragraphs]
    texts += [t for t in (p.text for p in doc.sections[0].header.paragraphs) if t]
    texts += [table.cell(0, 0).text]
    texts += [p.text for p in processing._iter_textbox_paragraphs(doc.part)]
    assert texts.count("VAL") == 4


def test_apply_conditionals_full(monkeypatch: Any) -> None:
    doc = Document()
    doc.add_paragraph("[[OPTION_1]]A[[/OPTION_1]][[OPTION_2]]B[[/OPTION_2]]")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "[[OPTION_1]]A[[/OPTION_1]][[OPTION_2]]B[[/OPTION_2]]"
    hdr_para = doc.sections[0].header.add_paragraph(
        "[[OPTION_1]]A[[/OPTION_1]][[OPTION_2]]B[[/OPTION_2]]"
    )
    hdr_table = doc.sections[0].header.add_table(rows=1, cols=1, width=Inches(1))
    hdr_table.cell(0, 0).text = "[[OPTION_1]]A[[/OPTION_1]][[OPTION_2]]B[[/OPTION_2]]"
    _patch_iter_textbox(monkeypatch, "[[OPTION_1]]A[[/OPTION_1]]")

    apply_conditionals(doc, {"{Action option}": "1"})

    assert [p.text for p in doc.paragraphs] == ["A"]
    assert table.cell(0, 0).text == "A"
    assert hdr_para.text == "A"
    assert [p.text for p in processing._iter_textbox_paragraphs(doc.part)] == ["A"]


@pytest.mark.parametrize(
    "choice,expected",
    [("1", "A"), ("2", "B"), ("3", "C"), ("4", "D")],
)
def test_apply_conditionals_golden_diff(
    tmp_path: Path, choice: str, expected: str
) -> None:
    """Applying conditionals should match the expected DOCX output for each option."""
    doc = Document()
    doc.add_paragraph(
        """
        [[OPTION_1]]A[[/OPTION_1]]
        [[OPTION_2]]B[[/OPTION_2]]
        [[OPTION_3]]C[[/OPTION_3]]
        [[OPTION_4]]D[[/OPTION_4]]
        """.strip()
    )
    processing.apply_conditionals(doc, {"{Action option}": choice})
    out_path = tmp_path / "out.docx"
    doc.save(str(out_path))

    expected_doc = Document()
    expected_doc.add_paragraph(expected)
    expected_path = tmp_path / "expected.docx"
    expected_doc.save(str(expected_path))

    # Compare document text rather than raw bytes to avoid differences in
    # metadata or ZIP structure.
    assert Document(str(out_path)).paragraphs[0].text == expected


def test_replace_placeholders_preserves_newlines() -> None:
    doc = Document()
    doc.add_paragraph("Start {P} End")
    replace_placeholders(doc, {"{P}": "Line1\nLine2"})
    text = doc.paragraphs[0].text
    assert "Line1" in text and "Line2" in text and "\n" in text


def test_replace_placeholders_retains_bold_style() -> None:
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("{name}")
    run.bold = True
    replace_placeholders(doc, {"{name}": "VAL"})
    assert para.runs[0].text == "VAL"
    assert para.runs[0].bold is True


def test_replace_placeholders_header_preserves_page_field() -> None:
    doc = Document()
    hdr = doc.sections[0].header
    p = hdr.add_paragraph()
    p.add_run("{x}")
    field_run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    field_run._r.append(fld)

    replace_placeholders(doc, {"{x}": "VAL"})
    assert "VAL" in p.text
    assert "PAGE" in p._p.xml


def test_replace_placeholders_textbox_preserves_formatting(monkeypatch: Any) -> None:
    doc = Document()
    para = Document().add_paragraph()
    run = para.add_run("{x}")
    run.italic = True
    monkeypatch.setattr(processing, "_iter_textbox_paragraphs", lambda part: [para])
    replace_placeholders(doc, {"{x}": "VAL"})
    para = processing._iter_textbox_paragraphs(doc.part)[0]
    assert para.runs[0].text == "VAL"
    assert para.runs[0].italic is True


def test_replace_placeholders_no_leftovers() -> None:
    doc = Document()
    doc.add_paragraph("{A} {B}")
    replace_placeholders(doc, {"{A}": "1", "{B}": "2"})
    text = "\n".join(p.text for p in doc.paragraphs)
    assert not re.search(r"\{.+?\}", text)
