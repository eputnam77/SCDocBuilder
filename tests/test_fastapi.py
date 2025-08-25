import pytest
from pathlib import Path
import pytest
from pathlib import Path
from docx import Document
from fastapi import UploadFile
import asyncio

# FastAPI requires the optional "python-multipart" or legacy "multipart"
# packages for form parsing.  Create lightweight stubs so the API module can be
# imported without these optional dependencies.
import types, sys

sys.modules.setdefault("python_multipart", types.ModuleType("python_multipart"))
sys.modules["python_multipart"].__version__ = "1.0.0"

sys.modules.setdefault("multipart", types.ModuleType("multipart"))
sys.modules.setdefault("multipart.multipart", types.ModuleType("multipart.multipart"))
sys.modules["multipart"].__version__ = "1.0.0"
sys.modules["multipart.multipart"].parse_options_header = lambda *a, **k: None

from scdocbuilder import api


def _make_docs(tmp_path: Path) -> tuple[Path, Path]:
    template = tmp_path / "t.docx"
    tdoc = Document()
    tdoc.add_paragraph("{Applicant name} {Airplane model}")
    tdoc.save(str(template))

    worksheet = tmp_path / "w.docx"
    wdoc = Document()
    wdoc.add_paragraph("Applicant name: Foo")
    wdoc.add_paragraph("Airplane model: Bar")
    wdoc.add_paragraph("Question 15:")
    wdoc.add_paragraph("Ans15")
    wdoc.add_paragraph("Question 16:")
    wdoc.add_paragraph("Ans16")
    wdoc.add_paragraph("Question 17:")
    wdoc.add_paragraph("Ans17")
    wdoc.save(str(worksheet))
    return template, worksheet


def _upload(path: Path) -> UploadFile:
    return UploadFile(path.open("rb"), filename=path.name)


def test_generate_endpoint_returns_doc(tmp_path: Path) -> None:
    template, worksheet = _make_docs(tmp_path)
    resp = asyncio.run(api.generate(_upload(template), _upload(worksheet)))
    assert resp.status_code == 200
    data = Path(resp.path).read_bytes()
    assert data.startswith(b"PK")


def test_generate_endpoint_returns_html(tmp_path: Path) -> None:
    template, worksheet = _make_docs(tmp_path)
    resp = asyncio.run(api.generate(_upload(template), _upload(worksheet), html=True))
    assert resp.status_code == 200
    assert b"<p" in resp.body


def test_health_endpoint_returns_ok() -> None:
    assert api.health() == {"status": "ok"}


def test_web_generate_returns_link(tmp_path: Path) -> None:
    template, worksheet = _make_docs(tmp_path)
    resp = asyncio.run(api.web_generate(_upload(template), _upload(worksheet)))
    assert resp.status_code == 200
    assert b"Download result" in resp.body
