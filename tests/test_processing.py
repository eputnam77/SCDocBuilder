from pathlib import Path
import typing
import pytest

if typing.TYPE_CHECKING:
    from docx import Document
else:
    pytest.importorskip("docx")
    from docx import Document

from scdocbuilder.processing import (
    extract_fields,
    replace_placeholders,
    apply_conditionals,
)


def _make_template(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Hello {name}")
    doc.save(str(path))


def test_extract_fields() -> None:
    doc = Document()
    doc.add_paragraph("Applicant name: Foo")
    fields = extract_fields(doc)
    assert fields["{Applicant name}"] == "Foo"


def test_replace_placeholders(tmp_path: Path) -> None:
    path = tmp_path / "t.docx"
    _make_template(path)
    doc = Document(str(path))
    replace_placeholders(doc, {"{name}": "Bar"})
    assert "Bar" in doc.paragraphs[0].text


def test_apply_conditionals() -> None:
    doc = Document()
    doc.add_paragraph("[[OPTION_1]]Yes[[/OPTION_1]][[OPTION_2]]No[[/OPTION_2]]")
    apply_conditionals(doc, {"{Action option}": "1"})
    assert doc.paragraphs[0].text == "Yes"


def test_apply_conditionals_across_runs() -> None:
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("[[OPTION_1]]Yes")
    para.add_run("[[/OPTION_1]]")
    para.add_run("[[OPTION_2]]No[[/OPTION_2]]")
    apply_conditionals(doc, {"{Action option}": "1"})
    assert doc.paragraphs[0].text == "Yes"


def test_apply_conditionals_four_options() -> None:
    doc = Document()
    doc.add_paragraph(
        """
        [[OPTION_1]]A[[/OPTION_1]]
        [[OPTION_2]]B[[/OPTION_2]]
        [[OPTION_3]]C[[/OPTION_3]]
        [[OPTION_4]]D[[/OPTION_4]]
        """.strip()
    )
    apply_conditionals(doc, {"{Action option}": "2"})
    assert doc.paragraphs[0].text == "B"


@pytest.mark.parametrize("choice, expected", [("10", "Ten"), ("2", "Two")])
def test_apply_conditionals_multi_digit_option(choice: str, expected: str) -> None:
    doc = Document()
    doc.add_paragraph("[[OPTION_10]]Ten[[/OPTION_10]][[OPTION_2]]Two[[/OPTION_2]]")
    apply_conditionals(doc, {"{Action option}": choice})
    assert doc.paragraphs[0].text == expected
