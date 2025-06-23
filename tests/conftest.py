# poetry run pytest

"""Test configuration and fixtures."""

import importlib
import pytest
from docx import Document

if importlib.util.find_spec("pytest_asyncio"):
    import pytest_asyncio  # type: ignore
    pytest_plugins = ["pytest_asyncio"]
    assert pytest_asyncio


def pytest_configure(config):
    config.addinivalue_line("markers", "asyncio: mark test as async/await test")


@pytest.fixture
def sample_worksheet():
    """Create a test document with basic fields."""
    doc = Document()
    doc.add_paragraph("Applicant name: Test Corp")
    doc.add_paragraph("Airplane manufacturer: Boeing")
    doc.add_paragraph("Airplane model: 787-TEST")
    doc.add_paragraph("14 CFR Part 25")
    return doc


@pytest.fixture
def multiline_worksheet():
    """Create a test document with multiline content."""
    doc = Document()
    # Fix the multiline field label to match what's expected
    doc.add_paragraph("Description:")  # This matches field mapping
    doc.add_paragraph("Line 1")
    doc.add_paragraph("Line 2")
    doc.add_paragraph("Line 3")
    # End marker
    doc.add_paragraph("Next field: value")
    return doc
