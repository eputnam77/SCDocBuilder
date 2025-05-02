# pytest tests/test_cli_integration.py -v

import os
import pytest
from pathlib import Path
from docx import Document
from faa_sc_filler.cli import parse_args, main

@pytest.fixture
def sample_files(tmp_path):
    """Create sample input files for testing."""
    template_path = tmp_path / "template.docx"
    worksheet_path = tmp_path / "worksheet.docx"
    output_path = tmp_path / "output.docx"
    
    # Create test template
    template = Document()
    template.add_paragraph("Contact: {SMEName}")
    template.save(template_path)
    
    # Create test worksheet
    worksheet = Document()
    worksheet.add_paragraph("Name of SME: John Doe")
    worksheet.save(worksheet_path)
    
    return template_path, worksheet_path, output_path

def test_cli_end_to_end(sample_files):
    """Test complete CLI workflow."""
    template_path, worksheet_path, output_path = sample_files

    # Run CLI with args
    main_result = main([
        "--template", str(template_path),
        "--worksheet", str(worksheet_path),
        "--output", str(output_path)
    ])
    
    # Verify results
    assert main_result == 0
    assert output_path.exists()
    
    # Check content was properly replaced
    output_doc = Document(output_path)
    assert "Contact: John Doe" in output_doc.paragraphs[0].text
    assert "[[NEED:" not in output_doc.paragraphs[0].text
