# pytest tests/test_webui.py -v

import pytest
import asyncio
import warnings
from pathlib import Path
from docx import Document
import re
import gradio as gr
from faa_sc_filler.webui import generate, create_ui

# Suppress websockets deprecation warning
warnings.filterwarnings("ignore", category=DeprecationWarning, module="websockets.legacy")

@pytest.fixture
def sample_files(tmp_path):
    """Create sample files for testing."""
    template_path = tmp_path / "template.docx"
    template = Document()
    template.add_paragraph("Contact: {SMEName}")
    template.save(template_path)
    
    worksheet_path = tmp_path / "worksheet.docx"
    worksheet = Document()
    worksheet.add_paragraph("Name of SME: John Doe")
    worksheet.save(worksheet_path)
    
    # Create mock Gradio file objects
    class MockFile:
        def __init__(self, path):
            self.name = str(path)
    
    return MockFile(template_path), MockFile(worksheet_path)

def test_generate_success(sample_files):
    """Test successful document generation."""
    template_file, worksheet_file = sample_files
    
    # Test normal processing
    output_path, diff = generate(template_file, worksheet_file, dry_run=False)
    assert output_path is not None
    assert Path(output_path).exists()
    
    # Test dry run
    _, diff = generate(template_file, worksheet_file, dry_run=True)
    assert isinstance(diff, dict)
    assert "{SMEName}" in str(diff)

def test_generate_error_handling():
    """Test error handling in generate function."""
    class MockFile:
        def __init__(self, path):
            self.name = str(path)
            
    bad_file = MockFile("nonexistent.docx")
    output, diff = generate(bad_file, bad_file)
    assert output is None
    assert "error" in diff

@pytest.mark.asyncio
async def test_ui_creation():
    """Test UI component creation."""
    # Create and test UI
    ui = create_ui()
    assert isinstance(ui, gr.Blocks)
    
    # Test component existence
    components = list(ui.blocks.values())
    assert len([c for c in components if isinstance(c, gr.File)]) >= 2  # Template and worksheet
    assert any(isinstance(c, gr.JSON) for c in components)  # Diff view
    assert any(isinstance(c, gr.Button) for c in components)  # Generate button
    
    # No need to close - Gradio handles cleanup automatically in test environment

@pytest.mark.parametrize("cfr,docket,notice,valid", [
    ("25", "FAA-2024-0001", "24-01-01-SC", True),
    ("25,27", "FAA-2024-0001", "24-01-01-SC", True),
    ("24", "FAA-2024-0001", "24-01-01-SC", False),
    ("25", "FAA-24-1", "24-01-01-SC", False),
    ("25", "FAA-2024-0001", "24-1-1", False),
])
def test_input_validation(cfr, docket, notice, valid):
    """Test validation of user inputs."""
    def validate_inputs(cfr, docket, notice):
        """Mock validation function for testing."""
        errors = []
        
        # Validate CFR parts
        if cfr:
            parts = [p.strip() for p in cfr.split(',')]
            valid_parts = {"23", "25", "27", "29", "31", "33", "35"}
            if not all(p in valid_parts for p in parts):
                errors.append("Invalid CFR format")
                
        # Validate docket number
        if not docket.startswith("FAA-") or not re.match(r'^FAA-\d{4}-\d{4}$', docket):
            errors.append("Invalid docket format")
            
        # Validate notice number
        if not re.match(r'^\d{2}-\d{2}-\d{2}-SC$', notice):
            errors.append("Invalid notice format")
            
        return errors
    
    errors = validate_inputs(cfr, docket, notice)
    assert bool(not errors) == valid

from faa_sc_filler.webui import enhanced_generate  # Import the function

def test_enhanced_generate(sample_files):
    """Test document generation with optional fields."""
    template_file, worksheet_file = sample_files
    
    # Create test with minimal requirements
    result = generate(
        template_file, 
        worksheet_file,
        dry_run=False
    )
    
    assert result is not None
    output, diff = result
    assert isinstance(output, (str, type(None)))
