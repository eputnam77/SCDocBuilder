import typing
from typing import Any, Mapping
import pytest
import builtins
import importlib.util

if typing.TYPE_CHECKING:
    from docx import Document
else:
    pytest.importorskip("docx")
    from docx import Document

from scdocbuilder.html_export import export_html

LIBS_AVAILABLE = (
    importlib.util.find_spec("mammoth") is not None
    and importlib.util.find_spec("bleach") is not None
)


def test_export_html_returns_string() -> None:
    doc = Document()
    doc.add_paragraph("Hello")

    html = export_html(doc)

    assert "<p" in html


def test_export_html_strips_script_tags() -> None:
    doc = Document()
    doc.add_paragraph("<script>alert('x')</script>")
    html = export_html(doc)
    assert "<script" not in html


def test_export_html_produces_heading_tags_fallback(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    original_import = builtins.__import__

    def fake_import(
        name: str,
        globals: Mapping[str, Any] | None = None,
        locals: Mapping[str, Any] | None = None,
        fromlist: tuple[str, ...] = (),
        level: int = 0,
    ) -> Any:
        if name in {"mammoth", "bleach"}:
            raise ImportError
        return original_import(name, globals, locals, fromlist, level)

    monkeypatch.setattr(builtins, "__import__", fake_import)

    doc = Document()
    doc.add_heading("Title", level=5)
    html = export_html(doc)
    assert "<h5>Title</h5>" in html


def test_export_html_fallback_paragraph(monkeypatch: pytest.MonkeyPatch) -> None:
    original_import = builtins.__import__

    def fake_import(name: str, *args: Any, **kwargs: Any) -> Any:
        if name in {"mammoth", "bleach"}:
            raise ImportError
        return original_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", fake_import)

    doc = Document()
    doc.add_paragraph("plain")
    html = export_html(doc)
    assert "<p>plain</p>" in html


def test_export_html_propagates_non_import_error(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Errors other than ImportError during import should bubble up."""
    original_import = builtins.__import__

    def fake_import(name: str, *args: Any, **kwargs: Any) -> Any:
        if name == "mammoth":
            raise RuntimeError("boom")
        return original_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", fake_import)

    doc = Document()
    doc.add_paragraph("plain")
    with pytest.raises(RuntimeError):
        export_html(doc)


def test_render_runs_formats_text() -> None:
    from scdocbuilder.html_export import _render_runs

    class Run:
        def __init__(self, text: str, bold: bool = False, italic: bool = False) -> None:
            self.text = text
            self.bold = bold
            self.italic = italic

    class Para:
        runs = [Run("a", bold=True), Run("b", italic=True)]

    assert _render_runs(Para()) == "<strong>a</strong><em>b</em>"


def test_heading_level_no_style() -> None:
    from scdocbuilder.html_export import _heading_level

    class Para:
        runs: list[Any] = []
        style = None

    assert _heading_level(Para()) == 0


@pytest.mark.skipif(not LIBS_AVAILABLE, reason="mammoth/bleach not installed")
def test_export_html_produces_heading_tags_with_libs() -> None:
    doc = Document()
    doc.add_heading("Title", level=5)
    html = export_html(doc)
    assert "<h5>Title</h5>" in html
