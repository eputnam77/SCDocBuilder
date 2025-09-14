import pytest
from pathlib import Path
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:
    from docx import Document
else:
    pytest.importorskip("docx")
    from docx import Document

from scdocbuilder.io import validate_input_files


def test_validate_input_files(tmp_path: Path) -> None:
    """validate_input_files should raise for bad paths and succeed for good."""
    t = tmp_path / "t.docx"
    w = tmp_path / "w.docx"
    Document().save(str(t))
    Document().save(str(w))

    validate_input_files(t, w)

    with pytest.raises(FileNotFoundError):
        validate_input_files(t, tmp_path / "missing.docx")


def test_validate_input_files_rejects_non_docx(tmp_path: Path) -> None:
    """Non-.docx inputs should raise ValueError."""
    t = tmp_path / "t.txt"
    w = tmp_path / "w.docx"
    t.write_bytes(b"\0")
    w.write_bytes(b"\0")
    with pytest.raises(ValueError):
        validate_input_files(t, w)


def test_validate_input_files_rejects_large_file(tmp_path: Path) -> None:
    """Files exceeding the size limit should be rejected."""
    big = tmp_path / "big.docx"
    big.write_bytes(b"0" * (11 * 1024 * 1024))
    with pytest.raises(ValueError):
        validate_input_files(big, big)


def test_validate_input_files_magic_error(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    """Errors from libmagic should not surface to callers."""
    t = tmp_path / "t.docx"
    w = tmp_path / "w.docx"
    Document().save(str(t))
    Document().save(str(w))

    import sys
    import types

    def fake_from_buffer(*args: Any, **kwargs: Any) -> str:
        raise OSError("missing magic database")

    fake_magic = types.SimpleNamespace(from_buffer=fake_from_buffer)
    monkeypatch.setitem(sys.modules, "magic", fake_magic)
    validate_input_files(t, w)
