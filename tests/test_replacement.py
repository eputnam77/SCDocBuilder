# pytest tests/test_replacement.py -v

from docx import Document
import logging

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def test_replacement():
    # Create a new document with a known placeholder
    doc = Document()
    doc.add_paragraph("Test SME Name: {SMEName}")
    doc.add_paragraph("Test Phone: {SMEPhone}")
    
    # Save the test template
    test_template = "test_template.docx"
    doc.save(test_template)
    logger.info(f"Created test template: {test_template}")
    
    # Now try to replace placeholders
    doc = Document(test_template)
    replacements = {
        "{SMEName}": "John Doe",
        "{SMEPhone}": "123-456-7890"
    }
    
    for paragraph in doc.paragraphs:
        logger.debug(f"Processing paragraph: {paragraph.text}")
        for run in paragraph.runs:
            logger.debug(f"Processing run: {run.text}")
            text = run.text
            for key, value in replacements.items():
                if key in text:
                    logger.info(f"Replacing {key} with {value}")
                    text = text.replace(key, value)
            if text != run.text:
                run.text = text
                logger.info(f"Updated text: {run.text}")
    
    # Save the result
    doc.save("test_result.docx")
    logger.info("Saved test result")

def test_multiple_replacements():
    doc = Document()
    doc.add_paragraph("SME: {SMEName}, Phone: {SMEPhone}, Email: {SMEEmail}")
    
    replacements = {
        "{SMEName}": "John Doe",
        "{SMEPhone}": "123-456-7890",
        "{SMEEmail}": "john@example.com"
    }
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text = run.text
            for key, value in replacements.items():
                if key in text:
                    text = text.replace(key, value)
            if text != run.text:
                run.text = text
    
    assert "John Doe" in doc.paragraphs[0].text
    assert "123-456-7890" in doc.paragraphs[0].text
    assert "john@example.com" in doc.paragraphs[0].text

def test_missing_replacement():
    doc = Document()
    doc.add_paragraph("Contact: {SMEName}")
    
    replacements = {}  # Empty replacements
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text = run.text
            for key, value in replacements.items():
                if key in text:
                    text = text.replace(key, value)
            if text != run.text:
                run.text = text
    
    assert "{SMEName}" in doc.paragraphs[0].text  # Placeholder should remain

def test_split_placeholder():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Contact: {")
    p.add_run("SMEName")
    p.add_run("}")
    
    replacements = {"{SMEName}": "John Doe"}
    
    # Get full paragraph text
    full_text = "".join(run.text for run in p.runs)
    if "{SMEName}" in full_text:
        # Replace in first run
        p.runs[0].text = full_text.replace("{SMEName}", "John Doe")
        # Clear other runs
        for run in p.runs[1:]:
            run.text = ""
    
    assert "John Doe" in p.text
    assert "{SMEName}" not in p.text
    assert len(p.text.strip()) > 0

def test_preserve_formatting():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("{SMEName}")
    run.bold = True
    run.italic = True
    
    replacements = {"{SMEName}": "John Doe"}
    
    text = run.text
    for key, value in replacements.items():
        if key in text:
            text = text.replace(key, value)
    if text != run.text:
        run.text = text
    
    assert "John Doe" in p.text
    assert p.runs[0].bold is True
    assert p.runs[0].italic is True

if __name__ == "__main__":
    test_replacement()
    test_multiple_replacements()
    test_missing_replacement()
    test_split_placeholder()
    test_preserve_formatting()
