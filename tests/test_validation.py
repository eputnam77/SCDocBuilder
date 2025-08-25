from pathlib import Path
import typing
import pytest

if typing.TYPE_CHECKING:
    from docx import Document
else:
    pytest.importorskip("docx")
    from docx import Document

from scdocbuilder import fill_template
from scdocbuilder import validation


def _make_template(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("{Applicant name} {Airplane model}")
    doc.save(str(path))


def test_missing_question_15_raises(tmp_path: Path) -> None:
    template = tmp_path / "t.docx"
    worksheet = tmp_path / "w.docx"
    _make_template(template)

    ws = Document()
    ws.add_paragraph("Applicant name: Foo")
    ws.add_paragraph("Airplane model: Bar")
    # Question 15 prompt without answer
    ws.add_paragraph("Question 15:")
    ws.save(str(worksheet))

    with pytest.raises(ValueError):
        fill_template(template, worksheet, tmp_path / "out.docx")


def test_next_question_not_used_as_answer(tmp_path: Path) -> None:
    """The validator should not treat the next question as the answer."""
    template = tmp_path / "t.docx"
    worksheet = tmp_path / "w.docx"
    _make_template(template)

    ws = Document()
    ws.add_paragraph("Applicant name: Foo")
    ws.add_paragraph("Airplane model: Bar")
    ws.add_paragraph("Question 15:")
    ws.add_paragraph("Question 16:")  # next question immediately
    ws.add_paragraph("Ans16")
    ws.add_paragraph("Question 17:")
    ws.add_paragraph("Ans17")
    ws.save(str(worksheet))

    with pytest.raises(ValueError):
        fill_template(template, worksheet, tmp_path / "out.docx")


def test_missing_question_triggers_error(tmp_path: Path) -> None:
    template = tmp_path / "t.docx"
    worksheet = tmp_path / "w.docx"
    _make_template(template)

    ws = Document()
    ws.add_paragraph("Applicant name: Foo")
    ws.add_paragraph("Airplane model: Bar")
    ws.add_paragraph("Question 15: Ans15")
    # Question 16 entirely absent
    ws.add_paragraph("Question 17: Ans17")
    ws.save(str(worksheet))

    with pytest.raises(ValueError):
        fill_template(template, worksheet, tmp_path / "out.docx")


def test_numbered_answer_is_accepted(tmp_path: Path) -> None:
    """Answers starting with enumerated bullets should not be treated as new questions."""
    ws = Document()
    ws.add_paragraph("Applicant name: Foo")
    ws.add_paragraph("Airplane model: Bar")
    ws.add_paragraph("Question 15:")
    ws.add_paragraph("1) Step one")
    ws.add_paragraph("Question 16: Ans16")
    ws.add_paragraph("Question 17: Ans17")
    # Should not raise
    validation.validate_mandatory_fields(ws)


def test_question_prompt_variants_are_detected() -> None:
    """Questions formatted as "15)" or "16-" should still be recognised."""
    ws = Document()
    ws.add_paragraph("Applicant name: Foo")
    ws.add_paragraph("Airplane model: Bar")
    ws.add_paragraph("15) Prompt")
    ws.add_paragraph("Ans15")
    ws.add_paragraph("16- Prompt")
    ws.add_paragraph("Ans16")
    ws.add_paragraph("Question 17:")
    ws.add_paragraph("Ans17")
    # Should not raise
    validation.validate_mandatory_fields(ws)
