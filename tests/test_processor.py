# pytest tests/test_processor.py -v

import pytest
from docx import Document
import sys
import os
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../src')))
from faa_sc_filler.processor import DocumentProcessor

def test_process_headers():
    # Create test doc with header placeholder
    doc = Document()
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = "Contact: {SMEName}"
    
    processor = DocumentProcessor()
    processor.replacements = {"{SMEName}": "John Doe"}
    processor.process_sections(doc)
    
    assert "John Doe" in section.header.paragraphs[0].text
    assert "{SMEName}" not in section.header.paragraphs[0].text

def test_conditional_blocks():
    doc = Document()
    doc.add_paragraph("[[OPTION_1]]Option 1 text[[/OPTION_1]]")
    doc.add_paragraph("[[OPTION_2]]Option 2 text[[/OPTION_2]]")
    
    processor = DocumentProcessor()
    processor.process_conditional_blocks(doc, selected_option=1)
    
    assert "Option 1 text" in doc.paragraphs[0].text
    assert len(doc.paragraphs) == 1  # Option 2 should be removed

def test_process_footer():
    doc = Document()
    section = doc.sections[0]
    footer = section.footer
    footer.paragraphs[0].text = "Page {PageNo}"
    
    processor = DocumentProcessor()
    processor.replacements = {"{PageNo}": "1"}
    processor.process_sections(doc)
    
    assert "1" in section.footer.paragraphs[0].text
    assert "{PageNo}" not in section.footer.paragraphs[0].text

def test_dry_run_processing():
    doc = Document()
    doc.add_paragraph("Hello {Name}")
    
    processor = DocumentProcessor()
    replacements = {"{Name}": "World"}
    
    _, diff = processor.process_document(doc, replacements, dry_run=True)
    assert diff["{Name}"] == {"old": "{Name}", "new": "World"}
    assert "Hello {Name}" in doc.paragraphs[0].text  # Original unchanged

def test_missing_field_handling():
    doc = Document()
    doc.add_paragraph("Contact: {SMEName}")
    
    processor = DocumentProcessor()
    processor.replacements = {}  # No replacement provided
    processor.process_document(doc, {})  # Pass document object directly
    
    assert "[[NEED:" in doc.paragraphs[0].text
    assert "SMEName" in doc.paragraphs[0].text

def test_multiple_replacements_same_paragraph():
    doc = Document()
    doc.add_paragraph("From: {SMEName}, Phone: {SMEPhone}")
    
    processor = DocumentProcessor()
    replacements = {
        "{SMEName}": "John Doe",
        "{SMEPhone}": "555-0123"
    }
    
    processor.process_document(doc, replacements)
    assert "From: John Doe, Phone: 555-0123" in doc.paragraphs[0].text

def test_table_cell_replacements():
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "{SMEName}"
    table.cell(0, 1).text = "{SMEPhone}"
    
    processor = DocumentProcessor()
    replacements = {
        "{SMEName}": "John Doe",
        "{SMEPhone}": "555-0123"
    }
    
    processor.process_document(doc, replacements)
    assert "John Doe" in table.cell(0, 0).text
    assert "555-0123" in table.cell(0, 1).text

def test_process_template_header():
    """Test processing of template header fields."""
    doc = Document()
    doc.add_paragraph("14 CFR Part {CFRPart}")
    doc.add_paragraph("[Docket No. {DocketNo}; Notice No. {NoticeNo}]")
    
    processor = DocumentProcessor()
    replacements = {
        "{CFRPart}": "25",
        "{DocketNo}": "FAA-2024-0001",
        "{NoticeNo}": "24-01-01-SC"
    }
    
    processor.process_document(doc, replacements)
    
    paragraphs = list(doc.paragraphs)
    assert "14 CFR Part 25" in paragraphs[0].text
    assert "FAA-2024-0001" in paragraphs[1].text
    assert "24-01-01-SC" in paragraphs[1].text

def test_summary_paragraph_processing():
    doc = Document()
    test_para = doc.add_paragraph("SUMMARY: Testing {SystemName} in {Location}")
    
    processor = DocumentProcessor()
    replacements = {
        "{SystemName}": "autopilot",  # lowercase input
        "{Location}": "aircraft"      # lowercase input
    }
    
    processor.process_document(doc, replacements)
    
    paragraph = doc.paragraphs[0]
    print("\nDebug Summary Paragraph:")
    print(f"Original paragraph text: {test_para.text}")
    print(f"Processed text: {paragraph.text}")
    print(f"Number of runs: {len(paragraph.runs)}")
    for i, run in enumerate(paragraph.runs):
        print(f"Run {i}: text='{run.text}', bold={run.bold}")
    
    # Verify exact content of each run
    assert len(paragraph.runs) == 2, "Should have exactly 2 runs"
    assert paragraph.runs[0].text == "SUMMARY: ", "First run should be 'SUMMARY: '"
    assert paragraph.runs[0].bold is True, "SUMMARY: should be bold"
    assert paragraph.runs[1].text == "Testing autopilot in aircraft", "Content should be in sentence case"
    assert paragraph.runs[1].bold is False, "Content should not be bold"

def test_invalid_replacements_type():
    doc = Document()
    doc.add_paragraph("Test {Field}")
    
    processor = DocumentProcessor()
    with pytest.raises(TypeError):
        processor.process_document(doc, "invalid")

def test_table_complex_cell_processing():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    # Fix: Combine into single paragraph with newline
    cell.text = "Name: {SMEName}\nPhone: {SMEPhone}"
    
    processor = DocumentProcessor()
    replacements = {
        "{SMEName}": "John Doe",
        "{SMEPhone}": "555-0123"
    }
    
    processor.process_document(doc, replacements)
    
    print(f"Cell text: {cell.text}")  # Debug output
    # Fix: Check complete cell text
    assert "Name: John Doe" in cell.text
    assert "Phone: 555-0123" in cell.text

def test_regex_pattern_handling():
    doc = Document()
    doc.add_paragraph("Testing {Complex.Field} and {Simple}")
    
    processor = DocumentProcessor()
    replacements = {
        "{Complex.Field}": "Complex Value",
        "{Simple}": "Simple Value"
    }
    
    processor.process_document(doc, replacements)
    
    assert "Testing Complex Value and Simple Value" in doc.paragraphs[0].text
