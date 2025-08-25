from pathlib import Path
from datetime import datetime, tzinfo
from typing import Any
import typing
import pytest

if typing.TYPE_CHECKING:
    from docx import Document
else:
    pytest.importorskip("docx")
    from docx import Document

from scdocbuilder import fill_template


def test_fill_template_writes_output(tmp_path: Path) -> None:
    template = tmp_path / "t.docx"
    worksheet = tmp_path / "w.docx"

    doc = Document()
    doc.add_paragraph("{Applicant name} {Airplane model}")
    doc.save(str(template))
    ws_doc = Document()
    ws_doc.add_paragraph("Applicant name: Foo")
    ws_doc.add_paragraph("Airplane model: Bar")
    ws_doc.add_paragraph("Question 15:")
    ws_doc.add_paragraph("Ans15")
    ws_doc.add_paragraph("Question 16:")
    ws_doc.add_paragraph("Ans16")
    ws_doc.add_paragraph("Question 17:")
    ws_doc.add_paragraph("Ans17")
    ws_doc.save(str(worksheet))

    output = tmp_path / "out.docx"
    result = fill_template(template, worksheet, output)

    assert result == output
    processed = Document(str(output))
    assert "Foo" in processed.paragraphs[0].text


def test_fill_template_uses_default_output(tmp_path: Path, monkeypatch: Any) -> None:
    template = tmp_path / "t.docx"
    worksheet = tmp_path / "w.docx"

    doc = Document()
    doc.add_paragraph("{Applicant name} {Airplane model}")
    doc.save(str(template))
    ws_doc = Document()
    ws_doc.add_paragraph("Applicant name: Foo")
    ws_doc.add_paragraph("Airplane model: Bar")
    ws_doc.add_paragraph("Question 15:")
    ws_doc.add_paragraph("Ans15")
    ws_doc.add_paragraph("Question 16:")
    ws_doc.add_paragraph("Ans16")
    ws_doc.add_paragraph("Question 17:")
    ws_doc.add_paragraph("Ans17")
    ws_doc.save(str(worksheet))

    fixed = datetime(2020, 1, 2, 3, 4, 5)

    class FakeDT(datetime):
        @classmethod
        def now(cls, tz: tzinfo | None = None) -> datetime:  # type: ignore[override]
            return fixed

    monkeypatch.setattr("scdocbuilder.datetime", FakeDT)

    result = fill_template(template, worksheet)
    expected = template.with_name("t_20200102_030405.docx")
    assert result == expected
    assert expected.exists()
