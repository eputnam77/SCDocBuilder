from typing import TYPE_CHECKING
from html.parser import HTMLParser

import pytest

if not TYPE_CHECKING:
    pytest.importorskip("docx")
from docx import Document

from scdocbuilder.html_export import export_html, _heading_level


def test_export_html_converts_headings_and_paragraphs() -> None:
    doc = Document()
    doc.add_heading("Title", level=1)
    doc.add_paragraph("Body text")
    html = export_html(doc)
    assert "<h1>" in html and "<p>Body text</p>" in html
    assert "style=" not in html


def test_export_html_sanitises_script_and_preserves_formatting() -> None:
    doc = Document()
    doc.add_paragraph("<script>alert('x')</script>")
    para = doc.add_paragraph()
    run = para.add_run("Italic")
    run.italic = True
    html = export_html(doc)
    assert "<script" not in html.lower()
    assert "<em>Italic</em>" in html


def test_export_html_tiptap_ready() -> None:
    doc = Document()
    doc.add_paragraph("Hello")
    html = export_html(doc)

    class Collector(HTMLParser):
        def __init__(self) -> None:
            super().__init__()
            self.tags: set[str] = set()

        def handle_starttag(
            self, tag: str, attrs: list[tuple[str, str | None]]
        ) -> None:
            self.tags.add(tag)

    parser = Collector()
    parser.feed(html)
    allowed = {"p", "h1", "h2", "h3", "ul", "ol", "li", "em", "strong", "a"}
    assert parser.tags <= allowed


def test_heading_level_handles_short_names() -> None:
    class Style:
        def __init__(self, name: str) -> None:
            self.name = name
            self.style_id = name

    class Paragraph:
        def __init__(self, style: Style) -> None:
            self.style = style

    para = Paragraph(Style("h2"))
    assert _heading_level(para) == 2
