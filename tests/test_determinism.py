import typing
from pathlib import Path

import pytest

if typing.TYPE_CHECKING:
    from docx import Document
else:
    pytest.importorskip("docx")
    from docx import Document

from scdocbuilder import fill_template


def _make_docs(tmp_path: Path) -> tuple[Path, Path]:
    template = tmp_path / "t.docx"
    tdoc = Document()
    tdoc.add_paragraph("{Applicant name} {Airplane model}")
    tdoc.save(str(template))

    worksheet = tmp_path / "w.docx"
    wdoc = Document()
    wdoc.add_paragraph("Applicant name: Foo")
    wdoc.add_paragraph("Airplane model: Bar")
    wdoc.add_paragraph("Question 15:")
    wdoc.add_paragraph("Ans15")
    wdoc.add_paragraph("Question 16:")
    wdoc.add_paragraph("Ans16")
    wdoc.add_paragraph("Question 17:")
    wdoc.add_paragraph("Ans17")
    wdoc.save(str(worksheet))
    return template, worksheet


def test_fill_template_repeatable(tmp_path: Path) -> None:
    template, worksheet = _make_docs(tmp_path)
    out1 = tmp_path / "out1.docx"
    out2 = tmp_path / "out2.docx"

    fill_template(template, worksheet, out1)
    fill_template(template, worksheet, out2)

    d1 = Document(str(out1))
    d2 = Document(str(out2))
    texts1 = [p.text for p in d1.paragraphs]
    texts2 = [p.text for p in d2.paragraphs]
    assert texts1 == texts2
